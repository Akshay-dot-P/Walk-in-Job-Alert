"""
resume_tailor.py — Research Framework Edition
==============================================
Generates ATS-optimised tailored DOCX+PDF resumes and measures how different
keyword strategies affect ATS scores and recruiter perception.

CHANGES IN THIS VERSION
A. Weighted JD text: title 3x, skills 2x, summary 1x for TF-IDF scoring
B. weak_detail threshold raised: ≥2 AMAZON_DETAIL_TOKENS required (was 1)
C. Domain-aware lens distribution: GRC/Risk/IAM/Forensics swap B2/B3 lens order
D. Research relevance gate: score < 3 sources filtered before LLM prompt
E. In-memory research cache keyed by (domain, job_title[:20])
F. sanitize_project_bullets: _has_purpose_clause check added
G. Dead code removed: AMAZON_LEGACY_COMPACT_FALLBACK_BULLETS,
   AMAZON_FALLBACK_BULLETS, get_weighted_amazon_fallback()
H. Parallel research: ThreadPoolExecutor(max_workers=3) for HTTP-bound
   experience + role-market research — Phase 1 keywords, Phase 2 research
   parallel, Phase 3 sequential generation+upload
I. Groq call budget counter: logs total calls at end of run
J. AI tailoring strategy: chooses projects, skills, work framing, and project
   bullet guidance from JD + grounded candidate/project evidence
K. Role-market intelligence: searches Reddit, GitHub, and web snippets for the
   target job title, then feeds market skills/keywords/project angles into AI
L. AppSec fallback profile + project/experience sanitizers for vague endings
   and cross-project keyword leakage

EXISTING FEATURES (unchanged)
B1. extract_keywords(jd_text) → semantic TF-IDF + cosine, Groq fallback
B2. SYNONYM_MAP + apply_synonyms() — safe post-generation expansion
B3. track_keyword_usage() — 2-3x coverage tracking
B4. dynamic_skills_augment() — JD keywords filtered via grounded evidence
B5. compute_metrics() → keyword_coverage, keyword_density, skills_count
B6. recruiter_simulate() → credibility, stuffing_suspicion, hireability
B7. enforce_single_page() — 5-tier STRICT single-page enforcement + page-fill
B8. Source-driven experience research (public resumes + Reddit)
B9. Three-lens career pivot framing
B10. generate_amazon_bullets_dynamic() — outcome-first, lens-framed
B11. jd_gap_analysis() — pre-pass gap table
B12. sanitize_project_bullets() — strips unverifiable metric claims

ADD TO requirements.txt:
  python-docx==1.1.2
  beautifulsoup4==4.12.3
  google-api-python-client==2.108.0
  pikepdf>=8.0
  pdfminer.six>=20221105
  scikit-learn>=1.3.0

WORKFLOW env:
  VALIDATION_MODE: normal            # lenient | normal | strict
  GROQ_GEN_MODEL: llama-3.3-70b-versatile
  GROQ_VAL_MODEL: llama-3.1-8b-instant
  GROQ_MAX_RETRIES: 1                # Fail fast on rate limits; fallbacks keep run moving
  GROQ_429_WAIT_BASE: 10
  GROQ_MIN_INTERVAL_SECONDS: 2
  GROQ_COOLDOWN_AFTER_429S: 3
  GROQ_COOLDOWN_SECONDS: 300
  MAX_JOBS_PER_RUN: 5
  AI_TAILORING: true                 # AI chooses projects/skills/work framing
  ROLE_MARKET_RESEARCH: true         # Looks up role-market skills/keywords
  ROLE_MARKET_LLM_SUMMARY: false     # Keep parallel research HTTP-only by default
  PROJECT_GITHUB_RESEARCH: true      # Pulls README evidence from project repos
  PROJECT_README_MAX_CHARS: 3500
  ALLOW_INFERRED_PROJECT_TOOLS: false # If true, can add JD tools without repo evidence
  RECRUITER_SIMULATION: false
  USE_LLM_SHORTENING: false
  ROLE_MARKET_MAX_WEB: 4
  ROLE_MARKET_MAX_GITHUB: 4
  ROLE_MARKET_MAX_REDDIT: 5
  PUBLIC_RESUME_RESEARCH: true
  WEB_RESUME_RESEARCH: true
  FORUM_RESEARCH: true
  RESUME_RESEARCH_MAX_WEB: 3
  FORUM_RESEARCH_MAX_POSTS: 6
  RESUME_RESEARCH_ALLOW_LINKEDIN: true
  RESUME_SOURCE_URLS: (newline or comma-separated PDF/DOCX URLs)
  REDDIT_RESEARCH_SUBS: (newline or comma-separated subreddits)
  CURRENT_ROLE_TITLE: Amazon operations associate - inventory reimbursement case triage
"""

from __future__ import annotations

import os, sys, re, json, time, io, base64, logging, requests, subprocess, tempfile, threading
from collections import Counter
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path
from urllib.parse import parse_qs, unquote, urlparse
from docx import Document
from docx.oxml.ns import qn
import gspread
from google.oauth2.service_account import Credentials
from bs4 import BeautifulSoup

# ─────────────────────────────────────────────────────────────────────────────
# Logging
# ─────────────────────────────────────────────────────────────────────────────
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
    handlers=[logging.StreamHandler(sys.stdout)],
)
logger = logging.getLogger(__name__)

# ─────────────────────────────────────────────────────────────────────────────
# Groq call budget counter  [NEW — tracks total API calls per run]
# ─────────────────────────────────────────────────────────────────────────────
_groq_call_count: int = 0
_groq_consecutive_429s: int = 0
_groq_cooldown_until: float = 0.0
_groq_last_call_ts: float = 0.0
_GROQ_LOCK = threading.Lock()

# ─────────────────────────────────────────────────────────────────────────────
# Config
# ─────────────────────────────────────────────────────────────────────────────
def _env_bool(name: str, default: bool = False) -> bool:
    raw = os.environ.get(name)
    if raw is None:
        return default
    return raw.strip().lower() not in {"0", "false", "no", "off", ""}


def _env_int(name: str, default: int, min_value: int = 0, max_value: int | None = None) -> int:
    try:
        value = int(os.environ.get(name, str(default)))
    except ValueError:
        value = default
    value = max(value, min_value)
    if max_value is not None:
        value = min(value, max_value)
    return value


def _env_float(name: str, default: float, min_value: float = 0.0,
               max_value: float | None = None) -> float:
    try:
        value = float(os.environ.get(name, str(default)))
    except ValueError:
        value = default
    value = max(value, min_value)
    if max_value is not None:
        value = min(value, max_value)
    return value


def _split_env_list(raw: str) -> list[str]:
    return [x.strip() for x in re.split(r"[\n,]+", raw or "") if x.strip()]


SHEET_NAME        = os.environ.get("SHEET_NAME", "WalkIn Jobs Bangalore")
GROQ_API_KEY      = os.environ.get("GROQ_API_KEY", "")
GROQ_GEN_MODEL    = os.environ.get("GROQ_GEN_MODEL", "llama-3.3-70b-versatile")
GROQ_VAL_MODEL    = os.environ.get("GROQ_VAL_MODEL", "llama-3.1-8b-instant")
GROQ_URL          = "https://api.groq.com/openai/v1/chat/completions"
GROQ_MAX_RETRIES  = _env_int("GROQ_MAX_RETRIES", 1, 1, 5)
GROQ_429_WAIT_BASE = _env_float("GROQ_429_WAIT_BASE", 10.0, 0.0, 120.0)
GROQ_MIN_INTERVAL_SECONDS = _env_float("GROQ_MIN_INTERVAL_SECONDS", 2.0, 0.0, 60.0)
GROQ_COOLDOWN_AFTER_429S = _env_int("GROQ_COOLDOWN_AFTER_429S", 3, 0, 20)
GROQ_COOLDOWN_SECONDS = _env_int("GROQ_COOLDOWN_SECONDS", 300, 0, 3600)
MAX_JOBS_PER_RUN  = _env_int("MAX_JOBS_PER_RUN", 5, 1, 20)
TEMPLATE_PATH     = Path(__file__).parent / "resume_template.docx"
GITHUB_TOKEN      = os.environ.get("GITHUB_TOKEN", "")
GITHUB_REPOSITORY = os.environ.get("GITHUB_REPOSITORY", "")
GITHUB_BRANCH     = os.environ.get("GITHUB_REF_NAME", "main")
RESUMES_FOLDER    = "resumes"
VALIDATION_MODE   = os.environ.get("VALIDATION_MODE", "normal").lower().strip()

PUBLIC_RESUME_RESEARCH         = _env_bool("PUBLIC_RESUME_RESEARCH", True)
WEB_RESUME_RESEARCH            = _env_bool("WEB_RESUME_RESEARCH", True)
FORUM_RESEARCH                 = _env_bool("FORUM_RESEARCH", True)
AI_TAILORING                   = _env_bool("AI_TAILORING", True)
ROLE_MARKET_RESEARCH           = _env_bool("ROLE_MARKET_RESEARCH", True)
ROLE_MARKET_LLM_SUMMARY        = _env_bool("ROLE_MARKET_LLM_SUMMARY", False)
PROJECT_GITHUB_RESEARCH        = _env_bool("PROJECT_GITHUB_RESEARCH", True)
ALLOW_INFERRED_PROJECT_TOOLS   = _env_bool("ALLOW_INFERRED_PROJECT_TOOLS", False)
RECRUITER_SIMULATION           = _env_bool("RECRUITER_SIMULATION", False)
USE_LLM_SHORTENING             = _env_bool("USE_LLM_SHORTENING", False)
RESUME_RESEARCH_MAX_WEB        = _env_int("RESUME_RESEARCH_MAX_WEB", 3, 0, 8)
FORUM_RESEARCH_MAX_POSTS       = _env_int("FORUM_RESEARCH_MAX_POSTS", 6, 0, 20)
PROJECT_README_MAX_CHARS       = _env_int("PROJECT_README_MAX_CHARS", 3500, 500, 8000)
ROLE_MARKET_MAX_WEB            = _env_int("ROLE_MARKET_MAX_WEB", 4, 0, 8)
ROLE_MARKET_MAX_GITHUB         = _env_int("ROLE_MARKET_MAX_GITHUB", 4, 0, 8)
ROLE_MARKET_MAX_REDDIT         = _env_int("ROLE_MARKET_MAX_REDDIT", 5, 0, 12)
RESUME_RESEARCH_ALLOW_LINKEDIN = _env_bool("RESUME_RESEARCH_ALLOW_LINKEDIN", True)
RESUME_SOURCE_URLS             = _split_env_list(os.environ.get("RESUME_SOURCE_URLS", ""))
REDDIT_RESEARCH_SUBS           = _split_env_list(os.environ.get(
    "REDDIT_RESEARCH_SUBS",
    "cybersecurityindia,cybersecurity,AskNetsec,cybersecurityjobs,resumes,developersIndia,cscareerquestionsIndia,IndianWorkplace",
))
CURRENT_ROLE_TITLE = os.environ.get(
    "CURRENT_ROLE_TITLE",
    "Amazon operations associate - inventory reimbursement case triage",
)

# [NEW] Research cache — keyed (domain:job_title[:20]), avoids redundant HTTP calls
# when multiple jobs share the same domain+title pattern within one run
_RESEARCH_CACHE: dict[str, dict] = {}
_PROJECT_EVIDENCE_CACHE: dict[str, str] = {}
_ROLE_MARKET_CACHE: dict[str, dict] = {}

SCOPES = [
    "https://www.googleapis.com/auth/spreadsheets",
    "https://www.googleapis.com/auth/drive",
]

# ─────────────────────────────────────────────────────────────────────────────
# SYNONYM MAP  (Feature 2)
# ─────────────────────────────────────────────────────────────────────────────
SYNONYM_MAP = {
    "ioc enrichment":          ["threat intelligence"],
    "log analysis":            ["SIEM monitoring"],
    "alert triage":            ["incident triage"],
    "threat detection":        ["anomaly detection"],
    "false positive analysis": ["alert tuning"],
    "incident escalation":     ["escalation workflows"],
    "spl correlation":         ["detection engineering"],
    "soar":                    ["security orchestration and automation"],
    "sigma rules":             ["detection-as-code"],
    "mitre att&ck":            ["TTP mapping"],
    "cvss severity":           ["vulnerability prioritisation"],
    "epss scoring":            ["exploit probability scoring"],
    "patch compliance":        ["remediation tracking"],
    "owasp top 10":            ["web application security"],
    "iam":                     ["identity and access management"],
    "cloudtrail":              ["cloud audit logging"],
    "guardduty":               ["cloud threat detection"],
    "cloud misconfiguration":  ["cloud security posture management"],
    "virustotal api":          ["threat intelligence feeds"],
    "osint enrichment":        ["open-source intelligence"],
    "typosquatting":           ["brand impersonation detection"],
    "audit documentation":     ["audit trail"],
    "root cause analysis":     ["investigative analysis"],
    "compliance monitoring":   ["regulatory compliance"],
    "nist csf":                ["cybersecurity framework"],
    "transaction monitoring":  ["financial crime detection"],
}


def apply_synonyms(text: str) -> str:
    if not text:
        return text
    applied = 0
    for term, aliases in SYNONYM_MAP.items():
        if applied >= 2:
            break
        alias = aliases[0]
        if alias.lower() in text.lower():
            continue
        pattern = re.compile(rf"(?<!\w){re.escape(term)}(?!\w)", re.IGNORECASE)
        def replacer(match):
            nonlocal applied
            if applied >= 2:
                return match.group(0)
            applied += 1
            return f"{match.group(0)} ({alias})"
        text, count = pattern.subn(replacer, text, count=1)
        if count > 0:
            continue
    return text


# ─────────────────────────────────────────────────────────────────────────────
# KEYWORD EXTRACTION  (Feature 1)
#
# HOW COSINE SIMILARITY WORKS HERE:
# ══════════════════════════════════
# 1. TF-IDF (Term Frequency-Inverse Document Frequency) builds a numeric
#    vector for each document.  Each dimension = one n-gram (1-word or 2-word
#    phrase).  The value for that dimension is high when the term appears
#    often in THIS document but rarely across all documents — i.e., it is
#    distinctive to this document, not generic.
#
# 2. We give TfidfVectorizer TWO documents:
#      doc[0] = the job description (what the employer wants)
#      doc[1] = _CANDIDATE_PROFILE (what the candidate actually has)
#    Result: a 2×N matrix where N = number of unique n-grams across both docs.
#
# 3. Element-wise multiplication:  scores = jd_arr * cand_arr
#    For each n-gram:
#      - jd_arr[i]   = TF-IDF weight of that term in the JD
#      - cand_arr[i] = TF-IDF weight of that term in the candidate profile
#    The product is 0 unless the term appears in BOTH documents.
#    High product = term is important in JD AND present in candidate's profile.
#    This is equivalent to the dot-product component of cosine similarity
#    (without the magnitude normalisation, which we don't need here because
#    we only want the ranked order, not the actual cosine angle).
#
# 4. We sort terms by score descending → top-15 are the keywords where
#    the JD's needs and the candidate's groundable experience overlap most.
#    Terms the JD uses but the candidate can't claim score near 0.
#    Terms the candidate has but the JD doesn't mention also score near 0.
#    Only the intersection surfaces — which is exactly what we want for ATS.
#
# 5. Fallback: if sklearn is not installed, deterministic catalog matching is
#    still used so keyword extraction does not burn Groq calls.
# ─────────────────────────────────────────────────────────────────────────────
_CYBER_STOPWORDS = {
    "experience", "knowledge", "understanding", "ability", "skill", "skills",
    "work", "working", "team", "role", "position", "candidate", "required",
    "preferred", "good", "strong", "excellent", "must", "will", "well",
    "including", "following", "responsible", "responsibilities", "etc",
    "years", "year", "day", "days", "time", "using", "used", "use",
    "help", "ensure", "support", "provide", "manage", "develop", "maintain",
    "cyber", "cybersecurity", "security", "information", "technical",
    "engineer", "analyst", "associate", "intern", "internship", "trainee",
    "summer", "hiring", "entry", "level", "solutions", "business",
}

_CANDIDATE_PROFILE = """
alert triage incident investigation log analysis threat detection escalation
false positive analysis root cause analysis audit documentation compliance tracking
policy enforcement anomaly detection pattern recognition evidence documentation
corrective actions seller claims reimbursement cases severity classification
splunk spl siem sigma rules soar python bash mitre attack ttp ioc virustotal
osint enrichment phishing typosquatting whois dns ssl abuseipdb urlscan
nessus openvas cvss epss nvd owasp sqli patch management cron api boto3 aws
iam cloudtrail guardduty cloud misconfiguration nist csf iso 27001 pci-dss
gdpr sox itgc risk assessment vendor risk transaction monitoring aml kyc
sanctions screening wireshark nmap tcp ip firewall ids ips endpoint security
windows linux active directory powershell cyber kill chain
"""

_KEYWORD_TOOL_PATTERNS = [
    ("Splunk", r"\bsplunk\b|\bspl\b"),
    ("SIEM", r"\bsiem\b"),
    ("Elastic SIEM", r"\belastic\b|\belk\b|\bkibana\b"),
    ("QRadar", r"\bqradar\b"),
    ("Microsoft Sentinel", r"\b(?:azure\s+)?sentinel\b|\bmicrosoft\s+sentinel\b"),
    ("CrowdStrike Falcon", r"\bcrowdstrike\b|\bfalcon\b"),
    ("Microsoft Defender", r"\bmicrosoft\s+defender\b|\bdefender\b|\bmde\b"),
    ("Sysmon", r"\bsysmon\b"),
    ("Sigma rules", r"\bsigma(?:\s+rules?)?\b"),
    ("SOAR", r"\bsoar\b|security orchestration"),
    ("VirusTotal", r"\bvirustotal\b"),
    ("AbuseIPDB", r"\babuseipdb\b"),
    ("URLScan.io", r"\burlscan(?:\.io)?\b"),
    ("Nessus", r"\bnessus\b"),
    ("OpenVAS", r"\bopenvas\b"),
    ("Burp Suite", r"\bburp(?:\s+suite)?\b"),
    ("OWASP ZAP", r"\bowasp\s+zap\b|\bzap\b"),
    ("Metasploit", r"\bmetasploit\b"),
    ("Nmap", r"\bnmap\b"),
    ("Wireshark", r"\bwireshark\b"),
    ("Python", r"\bpython\b"),
    ("Bash", r"\bbash\b"),
    ("PowerShell", r"\bpowershell\b"),
    ("SQL", r"\bsql\b"),
    ("AWS", r"\baws\b|amazon web services"),
    ("CloudTrail", r"\bcloudtrail\b"),
    ("GuardDuty", r"\bguardduty\b"),
    ("boto3", r"\bboto3\b"),
    ("CyberArk", r"\bcyberark\b"),
    ("SailPoint", r"\bsailpoint\b"),
    ("Okta", r"\bokta\b"),
    ("Active Directory", r"\bactive directory\b|\bad\b"),
    ("Linux", r"\blinux\b"),
    ("Windows", r"\bwindows\b"),
    ("Docker", r"\bdocker\b"),
    ("Kubernetes", r"\bkubernetes\b|\bk8s\b"),
    ("Trivy", r"\btrivy\b"),
    ("Semgrep", r"\bsemgrep\b"),
    ("Qualys", r"\bqualys\b"),
    ("Tenable", r"\btenable(?:\.io)?\b"),
]

_KEYWORD_ACTION_PATTERNS = [
    ("alert triage", r"\balert\s+triage\b|\btriag(?:e|ed|ing)\b"),
    ("incident investigation", r"\bincident\s+investigation\b|\binvestigat(?:e|ed|ing|ion)\b"),
    ("analysis", r"\blog\s+analysis\b|\banaly(?:s[ei]s|ze|zed|zing)\b"),
    ("security monitoring", r"\bmonitor(?:ing|ed)?\b"),
    ("threat detection", r"\bdetect(?:ion|ed|ing)?\b"),
    ("incident response", r"\brespond(?:ed|ing)?\b|\bresponse\b"),
    ("remediation tracking", r"\bremediat(?:e|ed|ing|ion)\b|\bpatch(?:ing)?\b"),
    ("risk assessment", r"\brisk\s+assessment\b|\bassess(?:ed|ing)?\b"),
    ("control testing", r"\bcontrol\s+test(?:ing)?\b"),
    ("security testing", r"\b(?:security|application|penetration)\s+test(?:ing)?\b"),
    ("audit documentation", r"\baudit(?:ed|ing)?\b|\bdocument(?:ed|ing|ation)?\b"),
    ("incident escalation", r"\bescalat(?:e|ed|ing|ion)\b"),
    ("IOC enrichment", r"\benrich(?:ed|ing|ment)?\b"),
    ("vulnerability scanning", r"\bscan(?:ned|ning)?\b"),
    ("threat hunting", r"\bhunt(?:ed|ing)?\b"),
    ("risk prioritization", r"\bprioriti[sz](?:e|ed|ing|ation)\b"),
    ("evidence review", r"\breview(?:ed|ing)?\b|\bvalidat(?:e|ed|ing|ion)\b"),
    ("automation", r"\bautomat(?:e|ed|ing|ion)\b"),
]

_KEYWORD_CONCEPT_PATTERNS = [
    ("SOC operations", r"\bsoc\b|security operations center|blue team"),
    ("incident response", r"\bincident response\b|\bdfir\b"),
    ("threat intelligence", r"\bthreat intelligence\b|\bcti\b"),
    ("IOC analysis", r"\bioc(?:s)?\b|indicator(?:s)? of compromise"),
    ("OSINT", r"\bosint\b|open source intelligence"),
    ("phishing analysis", r"\bphishing\b"),
    ("typosquatting", r"\btyposquat(?:ting)?\b|brand impersonat"),
    ("MITRE ATT&CK", r"\bmitre(?:\s+att&ck|\s+attack)?\b|\batt&ck\b"),
    ("TTP mapping", r"\bttps?\b|tactics techniques procedures"),
    ("Cyber Kill Chain", r"\bcyber kill chain\b"),
    ("vulnerability management", r"\bvulnerability management\b|\bpatch management\b"),
    ("vulnerability assessment", r"\bvulnerability assessment\b|\bvapt\b"),
    ("penetration testing", r"\bpenetration testing\b|\bpentest(?:ing)?\b"),
    ("application security", r"\bapplication security\b|\bappsec\b|product security"),
    ("OWASP Top 10", r"\bowasp(?:\s+top\s+10)?\b"),
    ("CVE analysis", r"\bcve(?:s)?\b|\bnvd\b"),
    ("CVSS scoring", r"\bcvss\b"),
    ("EPSS scoring", r"\bepss\b"),
    ("secure SDLC", r"\bsecure sdlc\b|\bsdlc\b|devsecops"),
    ("GRC", r"\bgrc\b|governance risk compliance"),
    ("compliance monitoring", r"\bcompliance\b|regulatory compliance"),
    ("control validation", r"\bcontrol(?:s)?\b|control validation"),
    ("audit evidence", r"\baudit evidence\b|evidence completeness|audit trail"),
    ("NIST CSF", r"\bnist(?:\s+csf)?\b"),
    ("ISO 27001", r"\biso\s*27001\b"),
    ("PCI-DSS", r"\bpci[-\s]?dss\b"),
    ("GDPR", r"\bgdpr\b"),
    ("SOX/ITGC", r"\bsox\b|\bitgc\b"),
    ("technology risk", r"\btechnology risk\b|\bit risk\b|cyber risk"),
    ("vendor risk", r"\bvendor risk\b|third[- ]party risk|tprm"),
    ("IAM", r"\biam\b|identity and access management"),
    ("identity governance", r"\bidentity governance\b|access governance|iga\b"),
    ("privileged access", r"\bprivileged access\b|\bpam\b"),
    ("cloud security", r"\bcloud security\b|aws security|azure security|gcp security"),
    ("cloud misconfiguration", r"\bcloud misconfiguration\b|\bcspm\b"),
    ("least privilege", r"\bleast privilege\b"),
    ("zero trust", r"\bzero trust\b"),
    ("network security", r"\bnetwork security\b"),
    ("IDS/IPS", r"\bids\b|\bips\b|intrusion detection|intrusion prevention"),
    ("firewall", r"\bfirewall\b"),
    ("endpoint security", r"\bendpoint security\b|\bedr\b|\bxdr\b"),
    ("packet analysis", r"\bpacket analysis\b|\bpcap\b|tcp/ip"),
    ("digital forensics", r"\bdigital forensics\b|\bforensics\b"),
    ("chain of custody", r"\bchain of custody\b"),
    ("fraud detection", r"\bfraud detection\b|\bfraud\b"),
    ("transaction monitoring", r"\btransaction monitoring\b"),
    ("AML/KYC", r"\baml\b|\bkyc\b|anti-money laundering|know your customer"),
    ("sanctions screening", r"\bsanctions screening\b|\bsanctions\b"),
    ("trust and safety", r"\btrust\s*(?:and|&)\s*safety\b"),
    ("privacy compliance", r"\bprivacy\b|data protection|dpdp|pdpb"),
]


def _normalise_keyword(term: str) -> str:
    return re.sub(r"\s+", " ", str(term or "")).strip(" .;:-_/|")


def _keyword_key(term: str) -> str:
    return _normalise_keyword(term).lower()


def _is_generic_keyword(term: str) -> bool:
    key = _keyword_key(term)
    if len(key) < 3:
        return True
    words = key.split()
    return all(w in _CYBER_STOPWORDS for w in words) or key in _CYBER_STOPWORDS


def _append_keyword(items: list[str], term: str, limit: int | None = None) -> None:
    term = _normalise_keyword(term)
    if not term or _is_generic_keyword(term):
        return
    term_key = _keyword_key(term)
    for i, existing in enumerate(list(items)):
        existing_key = _keyword_key(existing)
        if term_key == existing_key or term_key in existing_key:
            return
        if existing_key in term_key:
            items[i] = term
            return
    if limit is None or len(items) < limit:
        items.append(term)


def _regex_term_present(term: str, text: str) -> bool:
    pattern = re.escape(_normalise_keyword(term)).replace(r"\ ", r"\s+")
    return bool(re.search(rf"(?<!\w){pattern}(?!\w)", text or "", re.IGNORECASE))


def _catalog_keyword_matches(jd_text: str) -> list[tuple[str, str, int]]:
    matches = []
    seen = set()
    for category, patterns in (
        ("tools", _KEYWORD_TOOL_PATTERNS),
        ("actions", _KEYWORD_ACTION_PATTERNS),
        ("concepts", _KEYWORD_CONCEPT_PATTERNS),
    ):
        for label, pattern in patterns:
            match = re.search(pattern, jd_text or "", re.IGNORECASE)
            key = (category, _keyword_key(label))
            if match and key not in seen:
                seen.add(key)
                matches.append((category, label, match.start()))
    category_weight = {"tools": 0, "concepts": 1, "actions": 2}
    return sorted(matches, key=lambda item: (category_weight[item[0]], item[2], -len(item[1])))


def _keyword_category(term: str) -> str | None:
    term = _normalise_keyword(term)
    for category, patterns in (
        ("tools", _KEYWORD_TOOL_PATTERNS),
        ("concepts", _KEYWORD_CONCEPT_PATTERNS),
        ("actions", _KEYWORD_ACTION_PATTERNS),
    ):
        for label, pattern in patterns:
            if _keyword_key(term) == _keyword_key(label) or re.search(pattern, term, re.IGNORECASE):
                return category
    return None


def _keyword_grounding_text() -> str:
    parts = [_CANDIDATE_PROFILE]
    try:
        for project in PROJECTS.values():
            parts.append(project.get("title", ""))
            parts.extend(project.get("tech_base", []))
            parts.extend(project.get("bullets", []))
            for values in project.get("tech_swappable", {}).values():
                parts.extend(values)
    except NameError:
        pass
    return " ".join(parts).lower()


def _fallback_jd_terms(jd_text: str) -> list[str]:
    tokens = [
        t for t in re.findall(r"[a-zA-Z][a-zA-Z0-9+#.-]{2,}", (jd_text or "").lower())
        if t not in _CYBER_STOPWORDS
    ]
    counts = Counter(tokens)
    ranked = []
    for term, _ in counts.most_common(30):
        if _keyword_category(term) or _is_groundable_keyword(term):
            _append_keyword(ranked, term, 12)
    return ranked


def _extract_keywords_groq_fallback(jd_text: str) -> dict:
    system = "You are an ATS keyword analyst. Return ONLY valid JSON. No markdown."
    user = (
        f"Extract the top 10-15 most important keywords from this job description.\n"
        f"JD: {jd_text[:800]}\n\n"
        "Return raw JSON only:\n"
        '{"tools":["tool1","tool2"],'
        '"concepts":["concept1","concept2"],'
        '"actions":["action1","action2"],'
        '"ranked":["highest_priority",...up_to_15]}'
    )
    try:
        raw = _call_groq(system, user, GROQ_GEN_MODEL, max_tokens=300)
        return json.loads(_repair_json(raw))
    except Exception as exc:
        logger.warning("  Keyword extraction fallback failed: %s", exc)
        return {"tools": [], "concepts": [], "actions": [], "ranked": []}


def extract_keywords(jd_text: str) -> dict:
    """
    Hybrid JD keyword extraction:
    1. deterministic catalog match for concrete tools/actions/concepts
    2. TF-IDF ranking for remaining JD terms, with grounding as a bonus

    Accepts a pre-weighted JD string (title 3x, skills 2x, summary 1x).
    Falls back to deterministic catalog + token ranking if sklearn is absent.
    """
    if not jd_text or len(jd_text.strip()) < 30:
        return {"tools": [], "concepts": [], "actions": [], "ranked": []}

    tools, actions, concepts, ranked_terms = [], [], [], []
    catalog_matches = _catalog_keyword_matches(jd_text)
    for category, label, _pos in catalog_matches:
        if category == "tools":
            _append_keyword(tools, label, 8)
        elif category == "actions":
            _append_keyword(actions, label, 8)
        elif category == "concepts":
            _append_keyword(concepts, label, 8)
        _append_keyword(ranked_terms, label, 20)

    try:
        from sklearn.feature_extraction.text import TfidfVectorizer
        from sklearn.feature_extraction.text import ENGLISH_STOP_WORDS
        import numpy as np
    except ImportError:
        for term in _fallback_jd_terms(jd_text):
            category = _keyword_category(term)
            if category == "tools":
                _append_keyword(tools, term, 8)
            elif category == "actions":
                _append_keyword(actions, term, 8)
            elif category == "concepts":
                _append_keyword(concepts, term, 8)
            _append_keyword(ranked_terms, term, 20)
        logger.info(
            "  JD keywords (rule-based) — top 5: %s | tools: %s | actions: %s",
            ranked_terms[:5], tools[:3], actions[:3],
        )
        return {
            "tools":    tools[:6],
            "concepts": concepts[:6],
            "actions":  actions[:6],
            "ranked":   ranked_terms[:15],
        }

    vectorizer = TfidfVectorizer(
        ngram_range=(1, 3),
        max_features=450,
        stop_words=list(ENGLISH_STOP_WORDS.union(_CYBER_STOPWORDS)),
        token_pattern=r"(?u)\b[a-zA-Z][a-zA-Z0-9\+\#\-\.]{1,}\b",
        sublinear_tf=True,
    )
    grounding_text = _keyword_grounding_text()
    docs = [jd_text.lower(), grounding_text]
    try:
        tfidf_matrix = vectorizer.fit_transform(docs)
    except ValueError:
        return {
            "tools":    tools[:6],
            "concepts": concepts[:6],
            "actions":  actions[:6],
            "ranked":   ranked_terms[:15],
        }

    feature_names = vectorizer.get_feature_names_out()
    jd_arr   = tfidf_matrix[0].toarray()[0]
    cand_arr = tfidf_matrix[1].toarray()[0]

    scored_terms = []
    catalog_labels = {_keyword_key(label) for _cat, label, _pos in catalog_matches}
    for idx in np.argsort(jd_arr)[::-1]:
        term  = feature_names[idx]
        jd_score = jd_arr[idx]
        if jd_score < 0.001:
            break
        term = _normalise_keyword(term)
        words = term.split()
        if _is_generic_keyword(term) or any(w in _CYBER_STOPWORDS for w in words):
            continue
        if len(term) < 3:
            continue
        category = _keyword_category(term)
        grounded_bonus = min(cand_arr[idx] * 1.5, 0.7)
        category_bonus = {"tools": 0.55, "concepts": 0.40, "actions": 0.30}.get(category, 0.0)
        catalog_bonus = 0.8 if _keyword_key(term) in catalog_labels else 0.0
        phrase_bonus = 0.12 * min(len(words), 3)
        score = jd_score + grounded_bonus + category_bonus + catalog_bonus + phrase_bonus
        scored_terms.append((score, term, category))

    for _score, term, category in sorted(scored_terms, key=lambda item: item[0], reverse=True):
        if category == "tools":
            _append_keyword(tools, term, 8)
        elif category == "actions":
            _append_keyword(actions, term, 8)
        elif category == "concepts":
            _append_keyword(concepts, term, 8)
        _append_keyword(ranked_terms, term, 20)
        if len(ranked_terms) >= 18 and len(tools) >= 2 and len(actions) >= 2:
            break

    logger.info(
        "  JD keywords — top 5: %s | tools: %s | actions: %s | concepts: %s",
        ranked_terms[:5], tools[:3], actions[:3],
        concepts[:3],
    )
    return {
        "tools":    tools[:6],
        "concepts": concepts[:6],
        "actions":  actions[:6],
        "ranked":   ranked_terms[:15],
    }


# ─────────────────────────────────────────────────────────────────────────────
# KEYWORD INJECTION CONTROL  (Feature 3)
# ─────────────────────────────────────────────────────────────────────────────
def track_keyword_usage(content: dict, ranked_keywords: list) -> dict:
    bullet_keys = ["AMZ_B1","AMZ_B2","AMZ_B3","AMZ_B4","P1_B1","P1_B2","P1_B3","P2_B1","P2_B2","P2_B3"]
    all_text = " ".join(content.get(k, "") for k in bullet_keys)
    usage = {}
    for kw in ranked_keywords[:10]:
        pattern = re.compile(rf"\b{re.escape(kw)}\b", re.IGNORECASE)
        usage[kw] = len(pattern.findall(all_text))
    under   = [k for k, c in usage.items() if c == 0]
    over    = [k for k, c in usage.items() if c > 3]
    present = sum(1 for c in usage.values() if c > 0)
    logger.info(
        "  Keyword coverage: %d/%d present | under=%s over=%s",
        present, len(usage), under[:3], over[:2]
    )
    return usage


# ─────────────────────────────────────────────────────────────────────────────
# DYNAMIC SKILLS AUGMENTATION  (Feature 4)
# ─────────────────────────────────────────────────────────────────────────────
CANDIDATE_GROUNDABLE = {
    "splunk","spl","siem","sigma rules","soar","wireshark","nmap",
    "mitre att&ck","ttp","picerl","incident response","brute force detection",
    "lateral movement","privilege escalation","ioc","virustotal","telegram bot",
    "log analysis","alert triage","threat detection",
    "nessus","openvas","cve","cvss","epss","nvd","owasp","sqli",
    "patch management","remediation","bash scripting","cron","api",
    "phishing","osint","abuseipdb","urlscan","whois","dns","typosquatting",
    "threat intelligence","ioc enrichment","domain analysis",
    "iam","cloudtrail","guardduty","boto3","aws","s3","cloud security",
    "cloud misconfiguration","least privilege","cspm",
    "cloud security posture","cloud access controls","zero trust",
    "root cause analysis","audit documentation","escalation","triage",
    "policy enforcement","investigation","chain of custody",
    "nist csf","iso 27001","pci-dss","gdpr","sox","itgc",
    "compliance monitoring","risk assessment","vendor risk",
    "transaction monitoring","aml","kyc","sanctions screening",
    "tcp/ip","dns","http","firewall","ids","ips","endpoint security",
    "windows internals","linux","active directory","python","powershell",
    "cyber kill chain","osint enrichment","pcap",
}


def _is_groundable_keyword(keyword: str) -> bool:
    key = _keyword_key(keyword)
    if not key or _is_generic_keyword(key):
        return False
    groundable = {_keyword_key(item) for item in CANDIDATE_GROUNDABLE}
    if key in groundable:
        return True
    # Only phrase-level partial matching is safe. This avoids junk like
    # "domain" being accepted because "domain analysis" is groundable.
    if " " in key:
        return any(key in item or item in key for item in groundable)
    return False


def dynamic_skills_augment(profile_skills: dict, jd_keywords: dict) -> dict:
    ranked = jd_keywords.get("ranked", []) + jd_keywords.get("tools", [])
    if not ranked:
        return profile_skills
    skills = dict(profile_skills)
    safe   = []
    for kw in ranked[:15]:
        kl = kw.lower()
        if _is_groundable_keyword(kw):
            if not any(kl in v.lower() for v in skills.values()):
                safe.append(kw)
    if safe:
        existing  = skills.get("SK_V5", "")
        additions = ", ".join(safe[:3])
        skills["SK_V5"] = f"{existing}, {additions}" if existing else additions
        logger.info("  Dynamic skills +%s", additions)
    return skills


# ─────────────────────────────────────────────────────────────────────────────
# METRICS  (Feature 5)
# ─────────────────────────────────────────────────────────────────────────────
def compute_metrics(content: dict, jd_keywords: dict, ats_score) -> dict:
    ranked  = jd_keywords.get("ranked", [])
    bullets = [content.get(k, "") for k in
               ["AMZ_B1","AMZ_B2","AMZ_B3","AMZ_B4","P1_B1","P1_B2","P1_B3","P2_B1","P2_B2","P2_B3"]]
    all_text = " ".join(bullets).lower()
    hits = 0
    coverage = 0
    if ranked:
        hits = sum(
            1 for kw in ranked[:10]
            if re.search(rf"(?<!\w){re.escape(kw)}(?!\w)", all_text, re.IGNORECASE)
        )
        coverage = round(hits / min(len(ranked), 10) * 100)
    nonempty = [b for b in bullets if b.strip()]
    density  = 0.0
    if nonempty and ranked:
        total = sum(sum(1 for kw in ranked[:10] if kw.lower() in b.lower()) for b in nonempty)
        density = round(total / len(nonempty), 2)
    skill_vals   = [content.get(f"SK_V{i}", "") for i in range(1, 6)]
    skills_count = sum(len([x for x in v.split(",") if x.strip()]) for v in skill_vals)
    return {
        "ats_score":          ats_score,
        "keyword_coverage":   f"{coverage}%",
        "keyword_density":    str(density),
        "total_skills_count": str(skills_count),
    }


# ─────────────────────────────────────────────────────────────────────────────
# RECRUITER SIMULATION  (Feature 6)
# ─────────────────────────────────────────────────────────────────────────────
def recruiter_simulate(content: dict, job: dict) -> dict:
    bullets = "\n".join(
        f"• {content.get(k, '')}" for k in
        ["AMZ_B1","AMZ_B2","AMZ_B3","AMZ_B4","P1_B1","P1_B2","P1_B3","P2_B1","P2_B2","P2_B3"]
        if content.get(k)
    )
    skills = " | ".join(content.get(f"SK_V{i}", "") for i in range(1, 6))
    system = "You are an experienced India cybersecurity recruiter. Be direct. Return ONLY valid JSON."
    user   = (
        f"Role: {job['job_title']} at {job['company']}\n"
        f"Candidate: MCA grad, 1.5yr Amazon operations, 0 professional security experience.\n"
        f"Resume bullets:\n{bullets[:800]}\nSkills: {skills[:300]}\n\n"
        "Rate honestly:\n"
        '{"credibility":<1-10>,"stuffing_suspicion":<1-10>,"hireability":<1-10>,'
        '"explanation":"<one sentence each dimension, max 200 chars total>"}'
    )
    try:
        raw  = _call_groq(system, user, GROQ_VAL_MODEL, max_tokens=200)
        data = json.loads(_repair_json(raw))
        logger.info("  Recruiter: credibility=%s stuffing=%s hireability=%s",
                    data.get("credibility"), data.get("stuffing_suspicion"), data.get("hireability"))
        return data
    except Exception as exc:
        logger.warning("  Recruiter sim failed: %s", exc)
        return {"credibility":"N/A","stuffing_suspicion":"N/A","hireability":"N/A","explanation":""}


# ─────────────────────────────────────────────────────────────────────────────
# SKILL PROFILES
# ─────────────────────────────────────────────────────────────────────────────
SKILL_PROFILES = {
    "soc_security": {
        "SK_L1":"SOC Operations",      "SK_V1":"Alert triage, incident investigation, log analysis, threat detection, escalation, false positive analysis",
        "SK_L2":"SIEM & Monitoring",   "SK_V2":"Splunk (SPL), Elastic SIEM (basic), Windows Event Logs, Sysmon, Wireshark",
        "SK_L3":"Threat Intelligence", "SK_V3":"MITRE ATT&CK, IOC analysis, VirusTotal, OSINT enrichment, Cyber Kill Chain",
        "SK_L4":"Systems & Networking","SK_V4":"Windows internals, Linux fundamentals, TCP/IP, DNS, HTTP/S, firewall and IDS/IPS concepts",
        "SK_L5":"Automation",          "SK_V5":"Python, Bash (basic), regular expressions",
    },
    "soc_security_cloud": {
        "SK_L1":"SOC Operations",      "SK_V1":"Alert triage, incident investigation, log analysis, threat detection, escalation, false positive analysis",
        "SK_L2":"SIEM & Monitoring",   "SK_V2":"Splunk (SPL), Elastic SIEM (basic), Windows Event Logs, Sysmon, Wireshark",
        "SK_L3":"Threat Intelligence", "SK_V3":"MITRE ATT&CK, IOC analysis, VirusTotal, OSINT enrichment, Cyber Kill Chain",
        "SK_L4":"Systems & Networking","SK_V4":"Windows internals, Linux fundamentals, TCP/IP, DNS, HTTP/S, IDS/IPS, AWS (IAM, CloudTrail, GuardDuty), cloud security posture",
        "SK_L5":"Automation",          "SK_V5":"Python, Bash (basic), boto3, regular expressions",
    },
    "networking_entry": {
        "SK_L1":"Networking",          "SK_V1":"TCP/IP, OSI model, DNS, HTTP/S, firewall concepts, IDS/IPS concepts",
        "SK_L2":"OS & Scripting",      "SK_V2":"Linux (grep, netstat, log analysis), Windows internals, Active Directory (basics), PowerShell, Python, Bash",
        "SK_L3":"SIEM & Tools",        "SK_V3":"Splunk (SPL), Wireshark, PCAP analysis, Windows Event Logs, Nmap",
        "SK_L4":"Security Operations", "SK_V4":"Alert triage, log analysis, security monitoring, threat detection, incident escalation, endpoint security",
        "SK_L5":"Frameworks",          "SK_V5":"MITRE ATT&CK, Incident Response (PICERL), OWASP Top 10",
    },
    "grc_risk_fraud": {
        "SK_L1":"GRC & Compliance",    "SK_V1":"NIST CSF, ISO 27001, PCI-DSS, GDPR/PDPB, SOX/ITGC, compliance monitoring",
        "SK_L2":"Risk & Audit",        "SK_V2":"Risk assessment, control testing, audit documentation, vendor risk, RCSA basics",
        "SK_L3":"Fraud & AML",         "SK_V3":"Transaction monitoring, AML typologies, KYC/CDD, sanctions screening",
        "SK_L4":"Systems & Tools",     "SK_V4":"Windows internals, Linux fundamentals, Python, Excel, SQL (basic), TCP/IP basics",
        "SK_L5":"Frameworks",          "SK_V5":"MITRE ATT&CK, OWASP Top 10, Incident Response (PICERL), audit trail documentation",
    },
    "appsec_security": {
        "SK_L1":"Application Security", "SK_V1":"OWASP Top 10, vulnerability assessment, injection detection, broken authentication, SSRF analysis",
        "SK_L2":"Testing Tools",        "SK_V2":"Nessus, OpenVAS, OWASP ZAP, CVSS/EPSS severity classification",
        "SK_L3":"Security Review",      "SK_V3":"CVE research, remediation tracking, vulnerability-to-fix documentation, secure code review basics",
        "SK_L4":"Systems & Scripting",  "SK_V4":"Python, Bash, NVD API, Linux fundamentals, TCP/IP, HTTP/S, DNS",
        "SK_L5":"Frameworks",           "SK_V5":"MITRE ATT&CK, NIST CSF, secure development lifecycle, patch compliance tracking",
    },
}

DOMAIN_SKILL_PROFILE = {
    "SOC":"soc_security","VAPT":"soc_security","AppSec":"appsec_security","Forensics":"soc_security",
    "CloudSec":"soc_security_cloud","IAM":"soc_security_cloud",
    "Network":"networking_entry",
    "GRC":"grc_risk_fraud","Risk":"grc_risk_fraud","Fraud-AML":"grc_risk_fraud",
    "General":"soc_security",
}


def compute_skills(domain: str) -> dict:
    return dict(SKILL_PROFILES.get(
        DOMAIN_SKILL_PROFILE.get(domain, "soc_security"),
        SKILL_PROFILES["soc_security"]
    ))


# ─────────────────────────────────────────────────────────────────────────────
# PROJECTS
# ─────────────────────────────────────────────────────────────────────────────
PROJECTS = {
    "soc_auto": {
        "title": "SOC Automation and Threat Detection Lab",
        "github": "https://github.com/Akshay-dot-P/soc-threat-lab",
        "tech_base": ["Python","Splunk","Wireshark","Nmap","MITRE ATT&CK","Sigma rules"],
        "tech_swappable": {
            r"qradar|ibm qradar":                          ["QRadar"],
            r"elastic|kibana|elk":                         ["Elastic SIEM"],
            r"sentinel|azure sentinel|microsoft sentinel": ["Azure Sentinel"],
            r"crowdstrike|falcon|edr|xdr":                 ["CrowdStrike Falcon"],
            r"defender|microsoft defender|mde":            ["Microsoft Defender"],
            r"suricata|snort|zeek|ids\b|ips\b":            ["Suricata IDS"],
            r"burp suite|burp|appsec|web app":             ["Burp Suite"],
            r"metasploit|exploit|pentest|vapt":            ["Metasploit"],
            r"volatility|memory forensics|dfir":           ["Volatility"],
            r"soar|playbook|automation":                   ["SOAR playbook"],
            r"grafana|dashboard":                          ["Grafana"],
            r"sysmon|evtx":                                ["Sysmon"],
        },
        "bullets": [
            "Deployed Splunk SIEM with SPL correlation searches for brute-force detection, lateral movement, and privilege escalation; mapped TTPs to MITRE ATT&CK and wrote PICERL incident report.",
            "Built automated SOAR-style detection pipeline: Python script ingests Splunk alerts, runs IOC enrichment via VirusTotal API, and dispatches Telegram notifications with severity classification.",
            "Converted detection logic to Sigma rules (vendor-neutral format used by enterprise SOCs); performed TCP/IP analysis in Wireshark to detect SYN scans, DNS tunnelling, and plaintext credential exposure on unencrypted sessions.",
        ],
    },
    "vuln_scanner": {
        "title": "Vulnerability Scanner and Patch Prioritization Engine",
        "github": "https://github.com/Akshay-dot-P/vuln-scanner",
        "tech_base": ["Python","Bash","Nessus","OpenVAS","NVD API","CVSS/EPSS context"],
        "tech_swappable": {
            r"qualys":                            ["Qualys"],
            r"tenable":                           ["Tenable.io"],
            r"burp suite|burp|owasp|web app":     ["Burp Suite","OWASP ZAP"],
            r"nmap|network scan":                 ["Nmap"],
            r"epss|exploit probability":          ["EPSS API (FIRST.org)"],
            r"sast|bandit|semgrep|secure code":   ["Semgrep SAST"],
            r"container|docker|trivy|kubernetes": ["Trivy container scanner"],
        },
        "bullets": [
            "Built automated vulnerability assessment pipeline integrating Nessus and OpenVAS REST APIs in Python; generated CVE reports classified by CVSS severity and EPSS context from the FIRST.org API.",
            "Developed OWASP Top 10 automated web checker that sends crafted HTTP requests to detect injection, broken auth, and SSRF vulnerabilities; documented SQL injection exploit and parameterised query remediation.",
            "Automated scan scheduling via Bash and cron; built delta-scan logic to flag newly discovered CVEs and organize patch compliance tracking evidence.",
        ],
    },
    "phishing_osint": {
        "title": "Phishing and OSINT Threat Intelligence Tool",
        "github": "https://github.com/Akshay-dot-P/phishing-osint-tool",
        "tech_base": ["Python","VirusTotal API","AbuseIPDB","WHOIS","Telegram bot","DNS analysis"],
        "tech_swappable": {
            r"shodan|censys":                      ["Shodan API"],
            r"osint|open source intel|recon":      ["theHarvester"],
            r"phishing|url|domain|malicious link": ["URLScan.io"],
            r"fraud|aml|financial crime":          ["fraud pattern matching"],
            r"threat intel|cti|ioc|indicator":     ["MISP IOC feeds"],
            r"typosquat|brand|impersonat":         ["typosquatting detector"],
            r"email|spf|dkim|dmarc":               ["email header analyser"],
        },
        "bullets": [
            "Built multi-API threat intelligence pipeline: submits suspicious URLs/IPs to VirusTotal, AbuseIPDB, and URLScan.io; cross-references WHOIS registration age, DNS records, and SSL details for phishing analysis.",
            "Implemented typosquatting domain detector generating character-substitution variants of brand domains and checking live DNS resolution for brand-impersonation analysis.",
            "Deployed Telegram bot interface enabling analysts to submit URLs for live IOC enrichment; supports bulk CSV input/output for incident response workflows and includes OSINT enrichment via theHarvester for domain profiling.",
        ],
    },
}

DOMAIN_TO_PROJECTS = {
    "SOC":        ("soc_auto",       "phishing_osint"),
    "VAPT":       ("vuln_scanner",   "soc_auto"),
    "AppSec":     ("vuln_scanner",   "soc_auto"),
    "GRC":        ("phishing_osint", "vuln_scanner"),
    "Risk":       ("phishing_osint", "vuln_scanner"),
    "Fraud-AML":  ("phishing_osint", "vuln_scanner"),
    "CloudSec":   ("soc_auto",       "vuln_scanner"),
    "IAM":        ("soc_auto",       "phishing_osint"),
    "Forensics":  ("soc_auto",       "phishing_osint"),
    "Network":    ("soc_auto",       "vuln_scanner"),
    "General":    ("soc_auto",       "vuln_scanner"),
}

CONCEPT_SWAPPABLE = {
    "soc_auto": {
        r"grc|compliance|audit|iso\s*27001|nist|sox|itgc": [
            "SIEM-based compliance monitoring and security audit log retention",
            "automated security control validation via SPL correlation searches",
        ],
        r"fraud|aml|kyc|transaction.?monitor|financial.?crime": [
            "transaction anomaly detection via log correlation and pattern matching",
            "automated suspicious activity alerting with severity-based escalation",
        ],
        r"cloud|aws|azure|gcp|iam|saas": [
            "cloud security event monitoring and IAM access anomaly detection",
            "cross-account activity correlation for cloud-native threat detection",
        ],
        r"forensic|dfir|incident.?response|evidence|chain.?of.?custody": [
            "forensic-grade event timeline reconstruction from SIEM log artifacts",
            "automated evidence packaging with chain-of-custody documentation",
        ],
        r"network|ids|ips|firewall|packet|intrusion": [
            "network intrusion detection via deep packet analysis and IDS alert correlation",
            "protocol-level anomaly detection for network security monitoring",
        ],
    },
    "vuln_scanner": {
        r"grc|compliance|audit|iso\s*27001|nist|pci|sox": [
            "vulnerability findings mapped to compliance framework controls (PCI-DSS, NIST)",
            "audit-ready remediation tracking and compliance evidence organization",
        ],
        r"devsecops|appsec|ci/?cd|sdlc|secure.?cod|sast|dast": [
            "application security testing integrated with development release cycles",
            "vulnerability-to-remediation workflow for secure development lifecycle",
        ],
        r"cloud|aws|azure|container|docker|kubernetes": [
            "cloud infrastructure vulnerability assessment and misconfiguration detection",
            "continuous security scanning for cloud-deployed services and endpoints",
        ],
        r"fraud|aml|risk|financial": [
            "risk-informed vulnerability prioritization using CVE and EPSS context",
            "remediation tracking aligned with regulatory compliance workflows",
        ],
    },
    "phishing_osint": {
        r"grc|compliance|audit|vendor.?risk|third.?party|due.?diligence": [
            "domain reputation review for third-party vendor risk assessment",
            "risk evidence organization from multi-source OSINT intelligence",
        ],
        r"fraud|aml|kyc|transaction|financial.?crime|sanctions": [
            "KYC domain-verification workflow: WHOIS age, registrar, DNS, and SSL cross-check",
            "suspicious transaction indicator enrichment mapping domains to known fraud typologies",
        ],
        r"cti|threat.?intel|ioc|indicator|feed|hunt": [
            "IOC lifecycle management and multi-source threat intelligence correlation",
            "proactive infrastructure-based threat hunting via domain attribution analysis",
        ],
        r"risk|assessment|scoring": [
            "automated risk indicator enrichment for entity due diligence workflows",
            "domain and IP reputation review for risk documentation",
        ],
    },
}

BULLET_VARIANTS = {
    "soc_auto": {
        "cloud_iam": [
            "Deployed Splunk SIEM with SPL correlation searches to monitor IAM anomalies including unauthorized privilege escalation and suspicious cross-account access patterns; mapped cloud-relevant TTPs to MITRE ATT&CK and wrote PICERL incident report.",
            "Built automated cloud security detection pipeline: Python script ingests Splunk alerts for IAM policy violations, performs IOC enrichment via VirusTotal API, and dispatches Telegram notifications with severity classification.",
            "Developed Sigma-compatible detection rules for cloud-specific TTPs including credential abuse and lateral movement; performed network analysis in Wireshark to identify anomalous authentication and DNS traffic patterns in cloud environments.",
        ],
        "dfir_forensics": [
            "Deployed Splunk SIEM with SPL correlation searches for forensic event timeline reconstruction; tracked brute-force attempts, credential misuse, and script-based execution across host and network logs with MITRE ATT&CK TTP mapping.",
            "Built automated evidence collection pipeline: Python script ingests Splunk alerts, performs IOC enrichment via VirusTotal API, and generates severity-classified incident packages with chain-of-custody documentation for forensic investigation handoff.",
            "Converted detection logic to Sigma rules for cross-SIEM forensic portability; performed deep packet inspection in Wireshark to reconstruct attack sequences including SYN scans, DNS tunnelling, and credential exposure.",
        ],
        "network_ids": [
            "Deployed Splunk SIEM with SPL correlation searches for network intrusion detection; mapped brute-force, lateral movement, and privilege escalation alerts to MITRE ATT&CK with PICERL reporting.",
            "Built automated network alert triage pipeline: Python script ingests Splunk IDS alerts, performs IOC enrichment via VirusTotal API, and dispatches Telegram notifications with severity classification.",
            "Wrote Sigma rules (vendor-neutral IDS detection format) for enterprise network security; performed TCP/IP deep packet analysis in Wireshark to detect SYN scans, DNS tunnelling, port sweeps, and plaintext credential exposure across network segments.",
        ],
    },
    "vuln_scanner": {
        "devsecops_appsec": [
            "Built automated application security testing pipeline integrating Nessus and OpenVAS APIs in Python; generated vulnerability reports classified by CVSS severity with EPSS context from the FIRST.org API.",
            "Developed OWASP Top 10 automated application security checker detecting injection, broken authentication, SSRF, and XSS vulnerabilities; documented SQL injection exploit-to-remediation workflow with parameterised query fixes.",
            "Automated security scan scheduling via Bash and cron integrated with development cycles; built delta-scan logic to flag newly introduced CVEs for secure development lifecycle tracking.",
        ],
        "cloud_security": [
            "Built automated cloud infrastructure vulnerability assessment pipeline using Nessus and OpenVAS APIs in Python; generated CVE reports classified by CVSS severity with EPSS context from the FIRST.org API.",
            "Developed automated security checker for cloud-hosted applications testing OWASP Top 10 vulnerabilities including injection, broken authentication, and SSRF; documented remediation workflows for cloud service misconfigurations and exposed endpoints.",
            "Automated vulnerability scan scheduling via Bash and cron for cloud security monitoring; built delta-scan logic to detect newly exposed CVEs and organize cloud compliance evidence.",
        ],
        "compliance_audit": [
            "Built automated vulnerability assessment pipeline integrating Nessus and OpenVAS APIs in Python; generated audit-ready CVE reports classified by CVSS severity with EPSS context from the FIRST.org API.",
            "Developed OWASP Top 10 automated compliance checker validating web application security controls against regulatory requirements; documented vulnerability-to-remediation audit trails including SQL injection evidence and parameterised query fixes.",
            "Automated compliance scan scheduling via Bash and cron; built delta-scan logic to track remediation progress and organize audit evidence for patch compliance review.",
        ],
    },
    "phishing_osint": {
        "grc_risk_audit": [
            "Built multi-source risk assessment pipeline: submits vendor domains and IPs to VirusTotal, AbuseIPDB, and URLScan.io; cross-references WHOIS registration age, DNS records, and SSL certificate details for third-party due diligence.",
            "Implemented domain reputation assessment tool generating typosquatting variants of monitored domains and checking live DNS resolution for vendor and partner ecosystem review.",
            "Deployed automated risk assessment interface via Telegram bot enabling analysts to submit domains for enrichment; supports bulk CSV input/output for vendor risk assessment workflows and includes OSINT enrichment via theHarvester for domain profiling.",
        ],
        "fraud_aml": [
            "Built multi-API fraud intelligence pipeline: submits suspicious domains and IPs to VirusTotal, AbuseIPDB, and URLScan.io; cross-references WHOIS registration age, DNS records, and SSL details for KYC domain-verification workflows.",
            "Implemented typosquatting domain detector generating character-substitution variants of legitimate business domains and checking live DNS resolution for financial fraud infrastructure review.",
            "Deployed Telegram bot interface for live suspicious entity enrichment supporting bulk CSV input/output for investigation workflows; includes OSINT enrichment via theHarvester for domain profiling to support suspicious transaction report documentation.",
        ],
        "cti_threat_intel": [
            "Built multi-API cyber threat intelligence pipeline: submits IOCs to VirusTotal, AbuseIPDB, and URLScan.io; cross-references WHOIS registration data, DNS records, and SSL certificate details for intelligence products.",
            "Implemented typosquatting domain detector generating character-substitution variants of tracked infrastructure and checking live DNS resolution for infrastructure-based threat hunting.",
            "Deployed Telegram bot interface for real-time IOC enrichment enabling analysts to process indicators at scale; supports bulk CSV input/output for threat intelligence workflows and includes OSINT enrichment via theHarvester for comprehensive domain attribution.",
        ],
    },
}

DOMAIN_BULLET_VARIANT = {
    "SOC":       {},
    "VAPT":      {},
    "AppSec":    {"vuln_scanner": "devsecops_appsec"},
    "GRC":       {"vuln_scanner": "compliance_audit", "phishing_osint": "grc_risk_audit"},
    "Risk":      {"vuln_scanner": "compliance_audit", "phishing_osint": "grc_risk_audit"},
    "Fraud-AML": {"phishing_osint": "fraud_aml", "vuln_scanner": "compliance_audit"},
    "CloudSec":  {"soc_auto": "cloud_iam", "vuln_scanner": "cloud_security"},
    "IAM":       {"soc_auto": "cloud_iam", "phishing_osint": "cti_threat_intel"},
    "Forensics": {"soc_auto": "dfir_forensics", "phishing_osint": "cti_threat_intel"},
    "Network":   {"soc_auto": "network_ids"},
    "General":   {},
}

# ─────────────────────────────────────────────────────────────────────────────
# AMAZON EXPERIENCE CONSTANTS
# ─────────────────────────────────────────────────────────────────────────────
AMAZON_BASE = [
    "Triaged 50+ weekly inventory reimbursement cases by severity and policy eligibility, following structured case triage, escalation, and decision workflows.",
    "Conducted root cause analysis on seller claims, identified policy violations and anomalous patterns, and escalated findings to senior reviewers.",
    "Maintained audit-ready case documentation recording investigation findings, decisions, evidence notes, and corrective actions.",
    "Spotted recurring fraud patterns across 200+ weekly cases and flagged them early, reducing repeat-issue investigation time before escalation.",
]

AMAZON_RAW_FACTS = AMAZON_BASE   # alias used in generate_amazon_bullets_dynamic

AMAZON_KEYS         = ["AMZ_B1", "AMZ_B2", "AMZ_B3", "AMZ_B4"]
AMAZON_ACTION_VERBS = (
    "Triaged", "Investigated", "Analyzed", "Detected", "Documented",
    "Conducted", "Maintained", "Spotted",
)
AMAZON_MAX_CHARS    = 230
AMAZON_MIN_CHARS    = 95
AMAZON_DETAIL_TOKENS = (
    "50+ weekly", "200+ weekly", "severity", "escalat", "root cause",
    "audit-ready", "corrective action", "evidence", "policy", "anomal",
    "risk", "reviewer",
)

# ─────────────────────────────────────────────────────────────────────────────
# PURPOSE CLAUSE DETECTION
# ─────────────────────────────────────────────────────────────────────────────
PURPOSE_CLAUSE_RE = re.compile(
    r"\b(to (?:improve|optimize|enhance|ensure|streamline|boost|strengthen|"
    r"increase|reduce|maximize|support|drive|achieve|facilitate|accelerate|"
    r"promote|enable|allow|help|assist|maintain)\b"
    r"|in order to\b"
    r"|for (?:better|improved|enhanced|optimal|greater|effective)\b"
    r"|so (?:as to|that (?:we|the team|the org))\b)",
    re.IGNORECASE,
)


def _has_purpose_clause(bullet: str) -> bool:
    if not bullet:
        return False
    # Only check tail — avoids false positives on mid-bullet usage
    tail = bullet[-80:]
    return bool(PURPOSE_CLAUSE_RE.search(tail))


def _strip_purpose_clause(bullet: str) -> str:
    tail_start = max(0, len(bullet) - 80)
    match = PURPOSE_CLAUSE_RE.search(bullet, tail_start)
    if match:
        truncated = bullet[:match.start()].rstrip(" ,;—–-").rstrip()
        if len(truncated) >= 60:
            return truncated
    return bullet


# ─────────────────────────────────────────────────────────────────────────────
# AMAZON ROLE FOCUS MAPPING
# ─────────────────────────────────────────────────────────────────────────────
AMAZON_ROLE_FOCUS = {
    "soc": "SOC ROLES: prioritize alert triage, incident investigation, escalation workflows, and pattern analysis.",
    "security_operations": "SECURITY OPERATIONS ROLES: prioritize security monitoring, escalation handling, investigation records, and case handoffs.",
    "cybersecurity_analyst": "CYBERSECURITY ANALYST ROLES: prioritize alert triage, anomaly review, root cause analysis, and evidence-based reporting.",
    "incident_response": "INCIDENT RESPONSE / DFIR ROLES: prioritize incident prioritization, root cause analysis, impact assessment, escalation, and handoff notes.",
    "threat_intel": "THREAT INTELLIGENCE / OSINT ROLES: prioritize suspicious indicators, pattern recognition, abuse trends, and intelligence documentation.",
    "vulnerability_management": "VULNERABILITY MANAGEMENT / APPSEC ROLES: prioritize severity, risk impact, remediation tracking, repeat issues, and evidence review.",
    "cloud_security": "CLOUD SECURITY ROLES: prioritize policy exceptions, access review, risk signals, escalation triggers, and audit-ready notes.",
    "iam": "IAM / ACCESS GOVERNANCE ROLES: prioritize eligibility validation, policy exceptions, access review mindset, control gaps, and evidence notes.",
    "dlp": "DLP ROLES: prioritize policy violations, sensitive-case review, anomaly identification, escalation triggers, and investigation records.",
    "network_security": "NETWORK SECURITY ROLES: prioritize anomaly signals, monitoring prioritization, escalation triggers, and investigation notes.",
    "grc": "GRC / COMPLIANCE ROLES: prioritize audit documentation, compliance tracking, control validation, policy exceptions, and evidence gaps.",
    "it_audit": "IT AUDIT / ITGC ROLES: prioritize audit trails, control adherence, evidence completeness, control weaknesses, and documentation.",
    "technology_risk": "TECHNOLOGY RISK ROLES: prioritize risk identification, root cause analysis, control gaps, risk tracking, and escalation priorities.",
    "tprm": "THIRD-PARTY / VENDOR RISK ROLES: prioritize due diligence, evidence review, documentation gaps, risk scoring, and review prioritization.",
    "privacy": "PRIVACY / DATA PROTECTION ROLES: prioritize data handling, policy risk, compliance evidence, control gaps, and traceability.",
    "data_governance": "DATA GOVERNANCE ROLES: prioritize data quality, completeness checks, policy exceptions, traceability, and control review.",
    "fraud": "FRAUD ROLES: prioritize fraud detection, anomaly identification, pattern recognition, suspicious behavior, and case investigation.",
    "aml_kyc": "AML / KYC ROLES: prioritize evidence review, suspicious activity indicators, transaction monitoring analysis, and investigation records.",
    "trust_safety": "TRUST AND SAFETY ROLES: prioritize policy enforcement, abuse patterns, user risk, evidence-based review, and escalation decisions.",
    "risk_operations": "RISK OPERATIONS ROLES: prioritize process risk, exception handling, escalation, repeat-issue reduction, and workflow tracking.",
    "content_risk": "CONTENT RISK ROLES: prioritize policy review, abuse trends, anomaly detection, escalation, and consistent enforcement decisions.",
    "credit_risk": "CREDIT RISK ROLES: prioritize policy exceptions, eligibility decisions, risk indicators, evidence review, and review prioritization.",
    "general": "GENERAL OPERATIONS ROLES: prioritize process efficiency, escalation handling, structured decision-making, and workflow optimization.",
}

AMAZON_DOMAIN_FOCUS = {
    "SOC": "soc", "VAPT": "vulnerability_management", "AppSec": "vulnerability_management",
    "CloudSec": "cloud_security", "IAM": "iam", "Forensics": "incident_response",
    "Network": "network_security", "GRC": "grc", "Risk": "technology_risk",
    "Fraud-AML": "fraud", "General": "general",
}

AMAZON_FOCUS_PATTERNS = [
    ("trust_safety",          r"\b(trust\s*(?:and|&)\s*safety|abuse|policy enforcement|user safety|platform safety)\b"),
    ("content_risk",          r"\b(content risk|content moderation|content safety|policy review|moderation analyst)\b"),
    ("aml_kyc",               r"\b(aml|anti-money laundering|kyc|cdd|edd|transaction monitoring|sanctions|financial crime|str analyst|cft)\b"),
    ("fraud",                 r"\b(fraud|chargeback|loss prevention|suspicious reimbursement|fraud operations)\b"),
    ("privacy",               r"\b(privacy|data protection|gdpr|dpdp|pdpb|dpo|consent management|privacy compliance)\b"),
    ("data_governance",       r"\b(data governance|data quality|data lineage|metadata|records governance)\b"),
    ("tprm",                  r"\b(third[- ]party risk|tprm|vendor risk|supplier risk|supply chain risk|due diligence)\b"),
    ("credit_risk",           r"\b(credit risk|credit analyst|loan|underwriting|collections|portfolio risk)\b"),
    ("risk_operations",       r"\b(operational risk|risk operations|ops risk|rcsa|loss event|process risk)\b"),
    ("it_audit",              r"\b(it audit|is audit|itgc|technology audit|internal audit|control testing|sox)\b"),
    ("technology_risk",       r"\b(technology risk|cyber risk|it risk|enterprise risk|erm|risk analyst)\b"),
    ("grc",                   r"\b(grc|compliance|iso\s*27001|nist|pci[- ]dss|regulatory compliance|control validation)\b"),
    ("dlp",                   r"\b(dlp|data loss prevention|information protection|data leakage)\b"),
    ("iam",                   r"\b(iam|identity|access governance|identity governance|pam|idam|sailpoint|okta|cyberark|privileged access)\b"),
    ("cloud_security",        r"\b(cloud security|aws security|azure security|gcp security|cspm|cloudtrail|guardduty|cloud iam)\b"),
    ("incident_response",     r"\b(incident response|incident responder|dfir|forensic|digital forensics|ediscovery)\b"),
    ("threat_intel",          r"\b(threat intelligence|cti|osint|threat hunting|ioc|indicator|dark web|threat research)\b"),
    ("vulnerability_management", r"\b(product security|bug bounty|vulnerability disclosure|security review|"
                                 r"threat model|api security|vulnerability|vapt|penetration|pentest|"
                                 r"appsec|application security|devsecops|sast|dast|patch management)\b"),
    ("network_security",      r"\b(network security|ids|ips|firewall|intrusion|packet|endpoint security)\b"),
    ("soc",                   r"\b(soc|siem|blue team|alert triage|security monitoring|security operations center|tier\s*[12]|l[12]\s+analyst)\b"),
    ("security_operations",   r"\b(security operations|detect and respond|security monitoring analyst)\b"),
    ("cybersecurity_analyst", r"\b(cybersecurity analyst|cyber security analyst|security analyst|information security|infosec|cyber analyst)\b"),
]

AMAZON_WEIGHTED_CONTEXT = {
    "soc": ("structured alert triage and escalation workflows", "incident review and escalation handoffs", "fraud and anomaly patterns"),
    "security_operations": ("security monitoring and escalation workflows", "security operations handoffs", "operational risk and anomaly patterns"),
    "cybersecurity_analyst": ("alert triage, anomaly review, and escalation workflows", "evidence-based security reporting", "risk and anomaly patterns"),
    "incident_response": ("incident prioritization and escalation workflows", "incident records and handoffs", "repeat incident indicators"),
    "threat_intel": ("suspicious indicator review and escalation workflows", "threat intelligence review notes", "abuse and suspicious indicator patterns"),
    "vulnerability_management": ("severity review and remediation tracking workflows", "remediation evidence tracking", "severity and repeat-issue patterns"),
    "cloud_security": ("access review, policy exception, and escalation workflows", "audit-ready security review notes", "access-risk and abuse patterns"),
    "iam": ("eligibility validation and access review workflows", "access governance review evidence", "eligibility exception and control-gap patterns"),
    "dlp": ("policy violation review and escalation workflows", "policy investigation records", "policy exception and data-risk patterns"),
    "network_security": ("monitoring prioritization and escalation workflows", "investigation records and handoffs", "anomaly signal patterns"),
    "grc": ("compliance review, risk triage, and escalation workflows", "compliance tracking and control review", "policy risk and control-gap patterns"),
    "it_audit": ("control review and evidence escalation workflows", "audit trail completeness", "control weakness and process exception patterns"),
    "technology_risk": ("risk identification and escalation workflows", "risk tracking and review", "policy risk and control-gap patterns"),
    "tprm": ("due diligence, risk scoring, and escalation workflows", "vendor risk evidence review", "documentation gap and vendor risk patterns"),
    "privacy": ("data handling, policy risk, and escalation workflows", "privacy compliance evidence", "privacy risk and control-gap patterns"),
    "data_governance": ("data quality, completeness, and escalation workflows", "data governance review traceability", "data gap and policy exception patterns"),
    "fraud": ("fraud detection and escalation workflows", "transaction case review", "fraud patterns"),
    "aml_kyc": ("KYC evidence review and escalation workflows", "AML investigation records", "suspicious activity and transaction patterns"),
    "trust_safety": ("policy enforcement, user risk, and escalation workflows", "policy enforcement records", "abuse and user-risk patterns"),
    "risk_operations": ("process risk, exception handling, and escalation workflows", "risk operations tracking", "process risk and repeat-issue patterns"),
    "content_risk": ("content risk review and escalation workflows", "policy enforcement records", "abuse and policy violation patterns"),
    "credit_risk": ("eligibility review, risk indicators, and escalation workflows", "credit risk evidence review", "policy exception and credit risk patterns"),
    "general": ("structured case triage, escalation, and decision workflows", "workflow review and handoffs", "fraud and repeat-issue patterns"),
}


def _build_outcome_fallbacks(focus_key: str) -> dict:
    """Outcome-first static fallback — no purpose-clause endings."""
    triage, documentation, patterns = AMAZON_WEIGHTED_CONTEXT.get(
        focus_key, AMAZON_WEIGHTED_CONTEXT["general"]
    )
    return {
        "AMZ_B1": (
            "Triaged 50+ weekly inventory reimbursement cases by severity and policy eligibility, "
            f"applying {triage}, with zero missed escalations across reviewed queues."
        ),
        "AMZ_B2": (
            "Conducted root cause analysis on seller claims, identifying policy violations and "
            "anomalous patterns; escalated findings to senior reviewers without re-investigation loops."
        ),
        "AMZ_B3": (
            "Maintained audit-ready case documentation, recording findings, decisions, corrective actions, "
            f"and evidence notes for {documentation}, with no documentation gaps flagged."
        ),
        "AMZ_B4": (
            f"Spotted recurring {patterns} across 200+ weekly cases and flagged them early, "
            "cutting repeat-issue investigation cycles before escalation."
        ),
    }


def _build_source_driven_experience_fallback(experience_research: dict | None = None) -> dict:
    bridges = (experience_research or {}).get("bridges", []) or [
        "case triage -> severity classification and escalation routing",
        "audit-ready case notes -> evidence documentation and decision trails",
        "seller-claim anomalies -> suspicious pattern review and investigation handoffs",
    ]
    bridge_1 = bridges[0].split("->")[-1].strip()
    bridge_2 = bridges[1].split("->")[-1].strip() if len(bridges) > 1 else "evidence documentation and decision trails"
    bridge_3 = bridges[2].split("->")[-1].strip() if len(bridges) > 2 else "pattern review and investigation handoffs"
    return {
        "AMZ_B1": (
            "Triaged 50+ weekly inventory reimbursement cases by severity and policy eligibility, "
            f"translating high-volume case review into {bridge_1}."
        ),
        "AMZ_B2": (
            "Conducted root cause analysis on seller claims, identifying policy violations and "
            f"anomalous patterns for {bridge_3}."
        ),
        "AMZ_B3": (
            "Maintained audit-ready case documentation with findings, decisions, evidence notes, "
            f"and corrective actions aligned to {bridge_2}."
        ),
        "AMZ_B4": (
            "Spotted recurring fraud patterns across 200+ weekly cases and flagged repeat issues early, "
            "reducing investigation cycles before escalation."
        ),
    }


# ─────────────────────────────────────────────────────────────────────────────
# THREE-LENS CAREER PIVOT FRAMING
#
# [CHANGE C] Domain-aware lens distribution:
#   - Audit-heavy domains (GRC, Risk, IAM, Forensics): B2=AUDIT TRAIL, B3=PATTERN DETECTION
#   - All other domains:                                B2=PATTERN DETECTION, B3=AUDIT TRAIL
# ─────────────────────────────────────────────────────────────────────────────
THREE_LENS_FRAMES = {
    "SOC":       ("TRIAGE DISCIPLINE",  "PATTERN DETECTION",
                  "alert triage, escalation logic, severity classification, incident prioritization, case handoff"),
    "VAPT":      ("PATTERN DETECTION",  "TRIAGE DISCIPLINE",
                  "vulnerability prioritization, severity scoring, risk-based triage, CVSS classification, remediation tracking"),
    "Network":   ("TRIAGE DISCIPLINE",  "PATTERN DETECTION",
                  "alert escalation, monitoring prioritization, anomaly signal triage, incident routing"),
    "GRC":       ("AUDIT TRAIL",        "PATTERN DETECTION",
                  "audit-ready documentation, control validation, compliance tracking, evidence completeness"),
    "Risk":      ("AUDIT TRAIL",        "PATTERN DETECTION",
                  "risk identification, control gap analysis, escalation priorities, policy-based risk triage"),
    "Fraud-AML": ("PATTERN DETECTION",  "AUDIT TRAIL",
                  "fraud typology recognition, suspicious activity indicators, anomaly detection, KYC evidence review"),
    "CloudSec":  ("TRIAGE DISCIPLINE",  "AUDIT TRAIL",
                  "access review discipline, policy exception triage, IAM risk signals, audit-ready access records"),
    "IAM":       ("AUDIT TRAIL",        "TRIAGE DISCIPLINE",
                  "eligibility validation, access review discipline, policy exception handling, governance documentation"),
    "Forensics": ("AUDIT TRAIL",        "PATTERN DETECTION",
                  "evidence packaging, incident timeline reconstruction, root cause isolation, PICERL handoff"),
    "AppSec":    ("PATTERN DETECTION",  "AUDIT TRAIL",
                  "vulnerability triage, severity scoring, OWASP classification, remediation tracking, security review evidence"),
    "General":   ("TRIAGE DISCIPLINE",  "AUDIT TRAIL",
                  "structured escalation, severity-based prioritization, audit documentation, pattern recognition"),
}

THREE_LENS_DESCRIPTIONS = {
    "TRIAGE DISCIPLINE": (
        "Frame as severity classification, routing logic, escalation discipline, and SLA adherence."
    ),
    "PATTERN DETECTION": (
        "Frame as anomaly detection at scale, pattern recognition, repeat-issue flagging, and early escalation."
    ),
    "AUDIT TRAIL": (
        "Frame as audit evidence creation: decision trails, evidence notes, corrective actions, and documentation completeness."
    ),
}

# [CHANGE C] Domains where B2 and B3 lens order are swapped
_AUDIT_HEAVY_DOMAINS = {"GRC", "Risk", "IAM", "Forensics"}


def build_three_lens_context(domain: str, jd_text: str) -> str:
    primary, secondary, vocab = THREE_LENS_FRAMES.get(domain, THREE_LENS_FRAMES["General"])
    primary_desc   = THREE_LENS_DESCRIPTIONS[primary]
    secondary_desc = THREE_LENS_DESCRIPTIONS[secondary]

    # [CHANGE C] Swap B2/B3 for audit-heavy domains
    if domain in _AUDIT_HEAVY_DOMAINS:
        b2_lens = "AUDIT TRAIL"
        b3_lens = "PATTERN DETECTION"
    else:
        b2_lens = "PATTERN DETECTION"
        b3_lens = "AUDIT TRAIL"

    return f"""
THREE-LENS CAREER PIVOT FRAMING:
PRIMARY LENS [{primary}]: {primary_desc}
SECONDARY LENS [{secondary}]: {secondary_desc}
Domain vocabulary to weave naturally (1-2 per bullet): {vocab}
LENS DISTRIBUTION (domain-specific):
- Bullet 1: {primary} lens
- Bullet 2: {b2_lens} lens
- Bullet 3: {b3_lens} lens
- Bullet 4: {secondary} lens
END RULE: every bullet must end with an outcome, metric, or result — never a purpose clause.
"""


def get_amazon_focus_key(job: dict) -> str:
    domain    = str(job.get("domain", "")).strip()
    role_text = " ".join(str(job.get(k, "")) for k in ("job_title", "summary", "skills")).lower()
    for focus_key, pattern in AMAZON_FOCUS_PATTERNS:
        if re.search(pattern, role_text):
            return focus_key
    return AMAZON_DOMAIN_FOCUS.get(domain, "general")


def get_amazon_role_focus(job: dict) -> str:
    return AMAZON_ROLE_FOCUS[get_amazon_focus_key(job)]


# ─────────────────────────────────────────────────────────────────────────────
# DYNAMIC AMAZON BULLET GENERATION
# ─────────────────────────────────────────────────────────────────────────────
def generate_amazon_bullets_dynamic(job: dict, jd_keywords: dict,
                                    experience_research: dict | None = None,
                                    tailoring_strategy: dict | None = None) -> dict:
    domain    = str(job.get("domain", "General")).strip()
    jd_text   = f"{job.get('skills', '')} {job.get('summary', '')} {job.get('job_title', '')}"
    ranked_kw = jd_keywords.get("ranked", [])
    research_context = format_experience_research_prompt(experience_research, for_work_experience=True)
    role_focus = get_amazon_role_focus(job)
    strategy_focus = (tailoring_strategy or {}).get("work_focus", "")
    role_market = (tailoring_strategy or {}).get("role_market", {}) or {}
    market_exp_keywords = ", ".join(role_market.get("experience_keywords", [])[:10])
    avoid_claiming = ", ".join(role_market.get("avoid_claiming", [])[:8])
    kw_hint = (
        f"\nTop JD keywords to weave in (1-2 per bullet): {', '.join(ranked_kw[:8])}\n"
        if ranked_kw else ""
    )
    facts_block = "\n".join(f"{i+1}. {fact}" for i, fact in enumerate(AMAZON_RAW_FACTS))
    system = (
        "You are a career-pivot resume specialist. Reframe operations experience for cybersecurity roles. "
        "Return ONLY valid JSON. Write and not &. Escape internal quotes."
    )
    user = f"""
Job: {job.get('job_title', 'Role')} at {job.get('company', 'Company')}
Domain: {domain}
JD skills: {jd_text[:500]}
{kw_hint}
	{research_context}
		TARGET WORK-EXPERIENCE STRATEGY:
		{role_focus}
		{strategy_focus}
		Market experience keywords to use only when grounded: {market_exp_keywords}
		Avoid unsupported market claims: {avoid_claiming}

		RAW EXPERIENCE FACTS (do not invent beyond these):
	{facts_block}

Task:
1. Compare benchmark resume/forum signals to the target role.
2. Compare the target role to the transition role: {CURRENT_ROLE_TITLE}.
3. Rewrite only the current-role experience facts as transferable bullets.

Generate 4 bullets in fact order.
Rules:
- max 230 chars
- starts with {', '.join(AMAZON_ACTION_VERBS)}
- include 1-2 transferable target-role terms only when grounded in the current-role facts
- end with outcome/metric/result (NEVER a purpose clause like "to improve X")
- do not mention Amazon operations
- keep 50+ weekly / 200+ weekly / senior reviewer details grounded
- never claim tools, certifications, platforms, or source-resume achievements

Return only:
{{"AMZ_B1":"...","AMZ_B2":"...","AMZ_B3":"...","AMZ_B4":"..."}}
"""
    try:
        raw     = _call_groq(system, user, GROQ_GEN_MODEL, max_tokens=600)
        bullets = json.loads(_repair_json(raw))
        validated = {}
        for key in AMAZON_KEYS:
            b = re.sub(r"\s+", " ", str(bullets.get(key, ""))).strip().replace(" & ", " and ")
            if _has_purpose_clause(b):
                b = _strip_purpose_clause(b)
            valid = (
                b
                and len(b) <= AMAZON_MAX_CHARS
                and b.startswith(AMAZON_ACTION_VERBS)
                and not _has_purpose_clause(b)
            )
            validated[key] = b if valid else None
        fallbacks = _build_source_driven_experience_fallback(experience_research)
        for key in AMAZON_KEYS:
            if not validated.get(key):
                validated[key] = fallbacks[key]
        logger.info("  Dynamic Amazon bullets generated (domain=%s)", domain)
        return validated
    except Exception as exc:
        logger.warning("  Dynamic Amazon generation failed: %s", exc)
        return _build_source_driven_experience_fallback(experience_research)


# ─────────────────────────────────────────────────────────────────────────────
# JD GAP ANALYSIS
# ─────────────────────────────────────────────────────────────────────────────
_CANDIDATE_CAN_MEET = {
    "triage","escalat","prioriti","case management","queue","severity","sla","workflow",
    "investig","root cause","anomaly","pattern","analysis","document","audit","evidence",
    "record","report","trail","compliance","policy","control","fraud","risk","exception",
    "violation","transaction","monitoring","splunk","siem","python","bash","nessus",
    "openvas","virustotal","osint","phishing","cvss","epss","owasp","mitre","sigma",
    "wireshark","nmap","aws","boto3",
}
_CANDIDATE_CANNOT_MEET = {
    "pentest","penetration test","exploit","metasploit","burp","malware analysis",
    "reverse engineer","assembly","fuzzing","red team","5 years","7 years","10 years",
    "senior","lead","manager","cissp","cisa","ceh","oscp","giac",
}


def jd_gap_analysis(jd_text: str, jd_keywords: dict) -> dict:
    jd_lower = jd_text.lower()
    ranked   = jd_keywords.get("ranked", [])
    can_frame, cannot_meet = [], []
    for kw in ranked[:12]:
        kl = kw.lower()
        if any(signal in kl for signal in _CANDIDATE_CANNOT_MEET):
            cannot_meet.append(kw)
        elif any(signal in kl for signal in _CANDIDATE_CAN_MEET):
            can_frame.append(kw)
    for hard in _CANDIDATE_CANNOT_MEET:
        if re.search(rf"\b{re.escape(hard)}\b", jd_lower) and hard not in cannot_meet:
            cannot_meet.append(hard)
    gap_instruction = ""
    if cannot_meet:
        gap_instruction = (
            f"GAPS (do not fake these): {', '.join(cannot_meet[:5])}. "
            "Skip them and frame around transferable strengths."
        )
    if can_frame:
        gap_instruction += (
            f"\nSTRONG FRAMES: {', '.join(can_frame[:8])}. "
            "Weight bullets toward these."
        )
    return {"can_frame": can_frame, "cannot_meet": cannot_meet, "gap_instruction": gap_instruction}


_VAGUE_OUTCOME_RE = re.compile(
    r"\b(?:ensuring|for improved|for better|for enhanced|for timely|"
    r"for efficient|for optimal|for greater)\s+\w+(?:\s+\w+)?\s*[.;]?\s*$",
    re.IGNORECASE,
)

_ARTIFACT_WORDS_RE = re.compile(
    r",\s*\b(security|management|resolution|accountability|transparency|"
    r"allocation|efficiency)\s*[.;]?\s*$",
    re.IGNORECASE,
)


# ─────────────────────────────────────────────────────────────────────────────
# SANITIZE AMAZON BULLETS
# [CHANGE B] weak_detail now requires ≥ 2 AMAZON_DETAIL_TOKENS (was ≥ 1)
# ─────────────────────────────────────────────────────────────────────────────
def sanitize_amazon_bullets(content: dict, job: dict,
                            experience_research: dict | None = None) -> dict:
    fallback = _build_source_driven_experience_fallback(experience_research)
    explicit_comparison = re.compile(
        r"\b(mirroring|similar to|akin to|central to|used in|required for)\b.*\b(role|roles|soc|security|audit)\b",
        re.IGNORECASE,
    )
    for key in AMAZON_KEYS:
        bullet = re.sub(r"\s+", " ", str(content.get(key, ""))).strip()
        bullet = bullet.replace(" & ", " and ")
        if _has_purpose_clause(bullet):
            stripped = _strip_purpose_clause(bullet)
            if len(stripped) >= 60 and not _has_purpose_clause(stripped):
                bullet = stripped
        lowered     = bullet.lower()
        weak_density = len(bullet) < AMAZON_MIN_CHARS
        # [CHANGE B] require at least 2 detail tokens to prevent thin bullets
        weak_detail = sum(1 for token in AMAZON_DETAIL_TOKENS if token in lowered) < 2
        weak_vague_ending = bool(_VAGUE_OUTCOME_RE.search(bullet))
        artifact_ending   = bool(_ARTIFACT_WORDS_RE.search(bullet))
        invalid = (
            not bullet
            or len(bullet) > AMAZON_MAX_CHARS
            or "amazon operations" in lowered
            or explicit_comparison.search(bullet) is not None
            or not bullet.startswith(AMAZON_ACTION_VERBS)
            or weak_density
            or weak_detail
            or _has_purpose_clause(bullet)
            or weak_vague_ending
            or artifact_ending
        )
        content[key] = fallback[key] if invalid else bullet
    return content


# ─────────────────────────────────────────────────────────────────────────────
# SANITIZE PROJECT BULLETS
# [CHANGE F] _has_purpose_clause check added for project bullets
# ─────────────────────────────────────────────────────────────────────────────
PROJECT_METRIC_PATTERNS = [
    (re.compile(r"\b(?:used|using)?\s*P[12]_TECH\s+tools?\s*(?:with|to|and)?\s*", re.IGNORECASE), ""),
    (re.compile(r"\bP[12]_TECH\b", re.IGNORECASE), "project tools"),
    (re.compile(r"\bP[12]\s+project\b", re.IGNORECASE), "project"),
    (re.compile(r"\s*\([^)]*(?:index=\*|stats\s+count\s+by\s+src_ip|T1110|T1078|T1059)[^)]*\)", re.IGNORECASE), ""),
    (re.compile(r"\bT(?:1110|1078|1059)\b(?:\s*,\s*T(?:1110|1078|1059)\b)*", re.IGNORECASE), ""),
    (re.compile(r"\s*[—–-]\s*(?:reducing|reduced|improving|improved|increasing|increased|enabling|providing|cutting)\b.*$", re.IGNORECASE), ""),
    (re.compile(r"\b(?:reducing|reduced|improving|improved|increasing|increased|cutting)\b[^.;]*[.;]?", re.IGNORECASE), ""),
    (re.compile(r"\b(?:mean time|MTTR|MTTD|SLA deadlines?|Critical\s*=\s*24\s*hrs?|High\s*=\s*7\s*days?|Medium\s*=\s*30\s*days?)\b[^.;]*[.;]?", re.IGNORECASE), ""),
    (re.compile(r"\b(?:probability|confidence|quantitative|quantified|risk scores?|phishing probability score|fraud probability scores?)\b", re.IGNORECASE), "analysis"),
    (re.compile(r"\b(?:rarely used by freshers|faster|before reaching threat feeds)\b", re.IGNORECASE), ""),
]


def sanitize_project_bullets(content: dict) -> dict:
    for key in ["P1_B1","P1_B2","P1_B3","P2_B1","P2_B2","P2_B3"]:
        bullet = re.sub(r"\s+", " ", str(content.get(key, ""))).strip()
        if not bullet:
            continue
        for pattern, replacement in PROJECT_METRIC_PATTERNS:
            bullet = pattern.sub(replacement, bullet)
        bullet = re.sub(r"\s+([.;,])", r"\1", bullet)
        bullet = re.sub(r"\b(?:and|with|for|to|by)\s*[.;,]*$", "", bullet, flags=re.IGNORECASE)
        bullet = re.sub(r"^(?:helped|using|used)\b\s*", "", bullet, flags=re.IGNORECASE)
        bullet = re.sub(r"\s{2,}", " ", bullet).strip(" ;,-")
        # [CHANGE F] strip purpose clause from project bullets too
        if _has_purpose_clause(bullet):
            stripped = _strip_purpose_clause(bullet)
            if len(stripped) >= 40 and not _has_purpose_clause(stripped):
                bullet = stripped
        if not bullet:
            bullet = " "
        if bullet and bullet[-1] not in ".!?":
            bullet += "."
        content[key] = bullet
    return content


_PROJECT_EXCLUSIVE_TERMS: dict[str, re.Pattern] = {
    "vuln_scanner": re.compile(
        r"\b(typosquat|telegram bot|urlscan\.io|abuseipdb|whois|brand impersonat)\b",
        re.IGNORECASE,
    ),
    "phishing_osint": re.compile(
        r"\b(epss|nessus|openvas|nvd api|cvss severity|delta.scan|patch compliance|"
        r"soar pipeline|sigma rules)\b",
        re.IGNORECASE,
    ),
    "soc_auto": re.compile(
        r"\b(sql injection|owasp top|sqli|cve report|patch scheduling|patch compliance)\b",
        re.IGNORECASE,
    ),
}


def strip_cross_project_terms(content: dict, p1_key: str, p2_key: str) -> dict:
    for prefix, project_key in (("P1", p1_key), ("P2", p2_key)):
        exclusive_re = _PROJECT_EXCLUSIVE_TERMS.get(project_key)
        if not exclusive_re:
            continue
        for bullet_key in (f"{prefix}_B1", f"{prefix}_B2", f"{prefix}_B3"):
            bullet = str(content.get(bullet_key, ""))
            if not bullet or not exclusive_re.search(bullet):
                continue
            cleaned = exclusive_re.sub("", bullet)
            cleaned = re.sub(r"\s+([.;,])", r"\1", cleaned)
            cleaned = re.sub(r"\s{2,}", " ", cleaned).strip(" ,;.")
            if len(cleaned) > 40:
                logger.warning(
                    "  Cross-project term stripped from %s (%s): %s",
                    bullet_key, project_key, bullet[:70],
                )
                content[bullet_key] = cleaned + ("." if cleaned[-1] not in ".!?" else "")
    return content


# ─────────────────────────────────────────────────────────────────────────────
# COMPANY INTELLIGENCE
# ─────────────────────────────────────────────────────────────────────────────
COMPANY_INTEL = {
    "wipro":         {"framing":"24x7 SOC shifts, SLA discipline, shift documentation.",        "keywords":["24x7 SOC","SLA adherence","shift documentation"]},
    "tcs":           {"framing":"ISO 27001 ISMS, VAPT, compliance delivery.",                   "keywords":["ISMS","ISO 27001","compliance audit"]},
    "infosys":       {"framing":"Multi-client delivery, documentation quality.",                 "keywords":["documentation quality","multi-client"]},
    "hcl":           {"framing":"Cloud-native security, AWS, detection engineering.",            "keywords":["cloud security","AWS security"]},
    "cognizant":     {"framing":"24x7 SOC, BFSI compliance, investigation rigour.",             "keywords":["SOC operations","BFSI security"]},
    "capgemini":     {"framing":"GRC consulting, cloud security, European clients.",             "keywords":["GRC","NIST"]},
    "deloitte":      {"framing":"GRC consulting, ITGC/SOX audits, client risk reports.",        "keywords":["cyber risk advisory","ITGC","SOX"]},
    "kpmg":          {"framing":"ITGC/IS audit. CISA valued. Control testing.",                 "keywords":["IT audit","ITGC","SOX"]},
    "pwc":           {"framing":"Cyber risk advisory. RBI, SEBI, GDPR, PDPB.",                 "keywords":["cyber risk","regulatory compliance","GDPR"]},
    "ey":            {"framing":"EY GDS IT audit and GRC delivery.",                            "keywords":["GRC","IT audit","ITGC"]},
    "jpmorgan":      {"framing":"Technology risk, Basel III, AML/KYC operations.",              "keywords":["technology risk","AML","operational risk"]},
    "goldman sachs": {"framing":"Internal tech audit, ITGC, control testing.",                  "keywords":["technology audit","ITGC","SOX"]},
    "deutsche bank": {"framing":"KYC, AML, information security.",                              "keywords":["KYC","AML","transaction monitoring"]},
    "citi":          {"framing":"Fraud detection, risk analytics, anomaly detection.",           "keywords":["fraud detection","risk analytics"]},
    "amazon":        {"framing":"LP lens: Dive Deep, Bias for Action, automation mindset.",     "keywords":["dive deep","automation","AWS"]},
    "google":        {"framing":"Technical depth, automation, systems thinking.",               "keywords":["security engineering","automation"]},
    "microsoft":     {"framing":"Azure, AD, Sentinel. Growth mindset.",                        "keywords":["Azure security","Active Directory","Zero Trust"]},
    "hdfc bank":     {"framing":"Fraud detection, AML, RBI compliance.",                       "keywords":["AML","RBI compliance","fraud analytics"]},
    "bajaj finserv": {"framing":"Fraud/risk operations, NBFC compliance.",                     "keywords":["fraud operations","IT risk"]},
}


def get_company_intel(company_raw: str) -> dict | None:
    name = re.sub(r"\s*\(.*?\)\s*$", "", company_raw).strip().lower()
    for key, intel in COMPANY_INTEL.items():
        if key in name or name in key:
            logger.info("  Company intel: %s", key)
            return intel
    return None


def select_tools(project_key: str, jd_text: str, max_tools: int = 5) -> list[str]:
    proj     = PROJECTS[project_key]
    jd_lower = jd_text.lower()
    base     = list(proj["tech_base"])
    extra    = []
    evidence_lower = get_project_evidence(project_key).lower()
    for pattern, tools in proj["tech_swappable"].items():
        if re.search(pattern, jd_lower):
            for t in tools:
                if not ALLOW_INFERRED_PROJECT_TOOLS and t.lower() not in evidence_lower:
                    continue
                if t not in base and t not in extra:
                    extra.append(t)
    return (base + extra)[:max_tools]


def select_concepts(project_key: str, jd_text: str, max_concepts: int = 3) -> list[str]:
    concept_map = CONCEPT_SWAPPABLE.get(project_key, {})
    jd_lower    = jd_text.lower()
    concepts    = []
    for pattern, phrases in concept_map.items():
        if re.search(pattern, jd_lower):
            for phrase in phrases:
                if phrase not in concepts:
                    concepts.append(phrase)
    return concepts[:max_concepts]


def get_project_bullets(project_key: str, domain: str) -> list[str]:
    variant_name = DOMAIN_BULLET_VARIANT.get(domain, {}).get(project_key)
    if variant_name:
        variants = BULLET_VARIANTS.get(project_key, {})
        if variant_name in variants:
            return variants[variant_name]
    return PROJECTS[project_key]["bullets"]


_PROJECT_SIGNALS: dict[str, dict[str, set[str]]] = {
    "soc_auto": {
        "strong": {
            "splunk", "siem", "sigma rules", "soar", "alert triage",
            "blue team", "detection engineering", "threat hunting",
            "security operations center", "mitre att&ck", "picerl",
            "soc analyst", "log correlation",
        },
        "weak": {
            "incident", "detection", "monitoring", "correlation",
            "log analysis", "brute force", "lateral movement",
            "edr", "playbook", "sigma", "spl",
        },
    },
    "vuln_scanner": {
        "strong": {
            "vulnerability management", "cvss", "epss", "nessus", "openvas",
            "penetration testing", "vapt", "appsec", "application security",
            "owasp", "sast", "dast", "devsecops", "patch management",
            "security testing", "bug bounty", "product security",
            "vulnerability assessment", "api security", "threat modeling",
        },
        "weak": {
            "cve", "nvd", "patch", "remediation", "scanning", "sqli",
            "injection", "trivy", "qualys", "tenable", "vulnerability",
        },
    },
    "phishing_osint": {
        "strong": {
            "osint", "threat intelligence", "ioc enrichment", "virustotal",
            "abuseipdb", "typosquatting", "dark web", "cyber threat intelligence",
            "phishing analysis", "transaction monitoring", "fraud detection",
            "cti analyst", "threat intel",
        },
        "weak": {
            "phishing", "ioc", "whois", "urlscan", "enrichment",
            "indicator", "sanctions", "kyc", "aml", "dns lookup",
        },
    },
}

_DOMAIN_PROJECT_BOOST: dict[str, dict[str, int]] = {
    "SOC":       {"soc_auto": 4},
    "VAPT":      {"vuln_scanner": 5},
    "AppSec":    {"vuln_scanner": 6},
    "GRC":       {"phishing_osint": 3},
    "Risk":      {"phishing_osint": 3, "vuln_scanner": 1},
    "Fraud-AML": {"phishing_osint": 6},
    "CloudSec":  {"soc_auto": 4},
    "IAM":       {"soc_auto": 3},
    "Forensics": {"soc_auto": 4},
    "Network":   {"soc_auto": 4},
    "General":   {"soc_auto": 2},
}

_PROJECT_SIGNAL_IGNORE = {
    "security", "cyber", "cybersecurity", "analyst", "engineer", "associate",
    "intern", "trainee", "application", "information", "technical",
}


# ─────────────────────────────────────────────────────────────────────────────
# AI TAILORING STRATEGY
# ─────────────────────────────────────────────────────────────────────────────
def _project_readme_urls(github_url: str) -> list[str]:
    parsed = urlparse(github_url or "")
    if parsed.netloc.lower() != "github.com":
        return []
    parts = parsed.path.strip("/").split("/")
    if len(parts) < 2:
        return []
    owner, repo = parts[:2]
    names = ("README.md", "readme.md", "README.markdown", "README.rst")
    urls = []
    for branch in ("main", "master"):
        for name in names:
            urls.append(f"https://raw.githubusercontent.com/{owner}/{repo}/{branch}/{name}")
    return urls


def _scrub_markdown_text(text: str, max_chars: int = PROJECT_README_MAX_CHARS) -> str:
    text = re.sub(r"```.*?```", " ", text or "", flags=re.DOTALL)
    text = re.sub(r"!\[[^\]]*\]\([^)]+\)", " ", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"^[#>*\-\s]+", "", text, flags=re.MULTILINE)
    text = re.sub(r"\|", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text[:max_chars]


def _fetch_project_readme(project_key: str) -> str:
    if not PROJECT_GITHUB_RESEARCH:
        return ""
    if project_key in _PROJECT_EVIDENCE_CACHE:
        return _PROJECT_EVIDENCE_CACHE[project_key]
    github_url = PROJECTS.get(project_key, {}).get("github", "")
    readme = ""
    for url in _project_readme_urls(github_url):
        try:
            resp = requests.get(url, headers=_HDRS, timeout=8)
            if resp.status_code == 404:
                continue
            resp.raise_for_status()
            if len(resp.text) > 80:
                readme = _scrub_markdown_text(resp.text)
                break
        except Exception as exc:
            logger.debug("Project README fetch failed for %s via %s: %s", project_key, url, exc)
    _PROJECT_EVIDENCE_CACHE[project_key] = readme
    return readme


def get_project_evidence(project_key: str, domain: str = "General") -> str:
    proj = PROJECTS[project_key]
    bullets = get_project_bullets(project_key, domain)
    readme = _fetch_project_readme(project_key)
    parts = [
        f"Title: {proj['title']}",
        f"GitHub: {proj.get('github', '')}",
        f"Verified tools: {', '.join(proj.get('tech_base', []))}",
        "Verified project facts:",
        *[f"- {b}" for b in bullets],
    ]
    if readme:
        parts.append(f"GitHub README evidence: {readme[:PROJECT_README_MAX_CHARS]}")
    return "\n".join(parts)


def _project_catalog_for_prompt(job: dict) -> str:
    chunks = []
    for key, proj in PROJECTS.items():
        evidence = get_project_evidence(key, job.get("domain", "General"))
        chunks.append(f"[{key}] {proj['title']}\n{evidence[:1400]}")
    return "\n\n".join(chunks)


def _jd_blob(job: dict, jd_keywords: dict) -> str:
    return " ".join(filter(None, [
        str(job.get("job_title", "")),
        str(job.get("company", "")),
        str(job.get("domain", "")),
        str(job.get("summary", "")),
        str(job.get("skills", "")),
        " ".join(jd_keywords.get("ranked", [])[:12]),
        " ".join(jd_keywords.get("tools", [])[:8]),
    ]))


def _project_match_score(project_key: str, job: dict, jd_keywords: dict) -> int:
    jd = _jd_blob(job, jd_keywords).lower()
    proj = PROJECTS[project_key]
    haystack = " ".join([
        proj["title"],
        " ".join(proj.get("tech_base", [])),
        " ".join(proj.get("bullets", [])),
        " ".join(sum(proj.get("tech_swappable", {}).values(), [])),
        " ".join(sum(CONCEPT_SWAPPABLE.get(project_key, {}).values(), [])),
    ]).lower()
    score = 0
    ranked_terms = [str(kw).lower() for kw in jd_keywords.get("ranked", [])[:14] if str(kw).strip()]
    tool_terms   = [str(kw).lower() for kw in jd_keywords.get("tools", [])[:8] if str(kw).strip()]
    signal_terms = set(ranked_terms + tool_terms)
    signal_tiers = _PROJECT_SIGNALS.get(project_key, {})
    for kw in signal_terms:
        if kw in _PROJECT_SIGNAL_IGNORE or len(kw) < 4:
            continue
        if any(kw in signal for signal in signal_tiers.get("strong", set())):
            score += 3
        elif any(kw in signal for signal in signal_tiers.get("weak", set())):
            score += 1
    for term in set(re.findall(r"[a-zA-Z][a-zA-Z0-9+#.-]{2,}", jd)):
        if term in _RESEARCH_STOPWORDS:
            continue
        if re.search(rf"(?<!\w){re.escape(term)}(?!\w)", haystack):
            score += 1
    for kw in jd_keywords.get("ranked", [])[:10]:
        if kw and kw.lower() in haystack:
            score += 3
    mapped = DOMAIN_TO_PROJECTS.get(job.get("domain", "General"), ())
    if project_key in mapped:
        score += 5
    score += _DOMAIN_PROJECT_BOOST.get(job.get("domain", "General"), {}).get(project_key, 0)
    return score


def _fallback_project_pair(job: dict, jd_keywords: dict) -> tuple[str, str]:
    scored = sorted(
        PROJECTS.keys(),
        key=lambda key: _project_match_score(key, job, jd_keywords),
        reverse=True,
    )
    mapped = DOMAIN_TO_PROJECTS.get(job.get("domain", "General"), ("soc_auto", "vuln_scanner"))
    pair = []
    for key in scored + list(mapped):
        if key in PROJECTS and key not in pair:
            pair.append(key)
        if len(pair) == 2:
            break
    while len(pair) < 2:
        for key in ("soc_auto", "vuln_scanner", "phishing_osint"):
            if key not in pair:
                pair.append(key)
                break
    return pair[0], pair[1]


def _candidate_evidence_blob(project_keys: list[str] | tuple[str, ...] = ()) -> str:
    chunks = [_CANDIDATE_PROFILE, " ".join(AMAZON_RAW_FACTS)]
    for key in project_keys:
        if key in PROJECTS:
            chunks.append(get_project_evidence(key))
    return " ".join(chunks).lower()


def _skill_item_grounded(item: str, evidence_blob: str) -> bool:
    item = re.sub(r"\s+", " ", item or "").strip(" .;:-")
    if not item:
        return False
    lower = item.lower()
    groundable = set(CANDIDATE_GROUNDABLE)
    for proj in PROJECTS.values():
        groundable.update(t.lower() for t in proj.get("tech_base", []))
    if lower in groundable or any(lower in g or g in lower for g in groundable):
        return True
    words = [w for w in re.findall(r"[a-zA-Z][a-zA-Z0-9+#.-]{2,}", lower) if w not in _RESEARCH_STOPWORDS]
    return bool(words) and all(w in evidence_blob for w in words[:3])


def _normalize_skill_profile(raw_skills: dict, fallback: dict, project_keys: tuple[str, str]) -> dict:
    if not isinstance(raw_skills, dict):
        raw_skills = {}
    evidence_blob = _candidate_evidence_blob(project_keys)
    normalized = {}
    for i in range(1, 6):
        label_key, value_key = f"SK_L{i}", f"SK_V{i}"
        label = re.sub(r"\s+", " ", str(raw_skills.get(label_key, fallback.get(label_key, "")))).strip()
        value = str(raw_skills.get(value_key, fallback.get(value_key, "")))
        label = label[:32] if label else fallback.get(label_key, f"Skills {i}")
        items = []
        for item in re.split(r"[,|;/]+", value):
            item = re.sub(r"\s+", " ", item).strip(" .;:-")
            if item and item.lower() not in {x.lower() for x in items} and _skill_item_grounded(item, evidence_blob):
                items.append(item)
            if len(items) >= 8:
                break
        if len(items) < 3:
            items = [x.strip() for x in fallback.get(value_key, "").split(",") if x.strip()][:8]
        normalized[label_key] = label
        normalized[value_key] = ", ".join(items)
    return normalized


def _keywords_with_market(jd_keywords: dict, role_market_intel: dict | None = None) -> dict:
    merged = dict(jd_keywords or {})
    ranked = []
    for term in (
        list((jd_keywords or {}).get("ranked", []))
        + list((role_market_intel or {}).get("target_skills", []))
        + list((role_market_intel or {}).get("project_keywords", []))
        + list((role_market_intel or {}).get("experience_keywords", []))
    ):
        term = re.sub(r"\s+", " ", str(term)).strip()
        if term and term.lower() not in {x.lower() for x in ranked}:
            ranked.append(term)
    merged["ranked"] = ranked[:18]
    return merged


def _default_tailoring_strategy(job: dict, jd_keywords: dict,
                                role_market_intel: dict | None = None) -> dict:
    p1_key, p2_key = _fallback_project_pair(job, jd_keywords)
    base_skills = dynamic_skills_augment(
        compute_skills(job.get("domain", "General")),
        _keywords_with_market(jd_keywords, role_market_intel),
    )
    focus_key = get_amazon_focus_key(job)
    market_project_terms = (role_market_intel or {}).get("project_keywords", [])
    return {
        "projects": [p1_key, p2_key],
        "skills": base_skills,
        "work_focus": AMAZON_ROLE_FOCUS.get(focus_key, AMAZON_ROLE_FOCUS["general"]),
        "project_guidance": {
            p1_key: select_concepts(p1_key, _jd_blob(job, jd_keywords) + " " + " ".join(market_project_terms)),
            p2_key: select_concepts(p2_key, _jd_blob(job, jd_keywords) + " " + " ".join(market_project_terms)),
        },
        "role_market": role_market_intel or {},
        "source": "fallback-score",
    }


def build_tailoring_strategy(job: dict, jd_keywords: dict,
                             experience_research: dict | None = None,
                             role_market_intel: dict | None = None) -> dict:
    fallback = _default_tailoring_strategy(job, jd_keywords, role_market_intel)
    if not AI_TAILORING or not GROQ_API_KEY:
        return fallback

    system = (
        "You are a resume tailoring strategist. Return ONLY valid JSON. "
        "Use only the candidate/project/current-role evidence provided. Do not invent tools, certs, employers, or metrics."
    )
    user = f"""
JOB:
Title: {job.get('job_title', '')}
Company: {job.get('company', '')}
Domain: {job.get('domain', '')}
Summary: {job.get('summary', '')}
Skills/JD terms: {job.get('skills', '')}
Top extracted keywords: {', '.join(jd_keywords.get('ranked', [])[:12])}

CURRENT ROLE FACTS:
{chr(10).join('- ' + fact for fact in AMAZON_RAW_FACTS)}

PROJECT CATALOG:
{_project_catalog_for_prompt(job)}

	ROLE RESEARCH SIGNALS:
	{(experience_research or {}).get('summary', '')[:1400]}

	ROLE-MARKET INTELLIGENCE FROM REDDIT/GITHUB/WEB:
	{(role_market_intel or {}).get('summary', '')[:1800]}
	Market target skills: {', '.join((role_market_intel or {}).get('target_skills', [])[:14])}
	Market experience keywords: {', '.join((role_market_intel or {}).get('experience_keywords', [])[:12])}
	Market project angles: {'; '.join((role_market_intel or {}).get('project_angles', [])[:8])}
	Unsupported / avoid claiming: {', '.join((role_market_intel or {}).get('avoid_claiming', [])[:10])}

	Task:
	1. Pick exactly 2 project keys from: {', '.join(PROJECTS.keys())}. Keep their titles unchanged later.
	2. Build a 5-row skills profile tailored to the JD + market intelligence using only grounded candidate/project evidence.
	3. Give short work-experience focus guidance using market experience keywords only when they fit current-role facts.
	4. Give 2-3 project guidance phrases per selected project using market project angles where supported.

Return raw JSON only:
{{
  "projects":["project_key_1","project_key_2"],
  "skills":{{"SK_L1":"...","SK_V1":"a, b, c","SK_L2":"...","SK_V2":"...","SK_L3":"...","SK_V3":"...","SK_L4":"...","SK_V4":"...","SK_L5":"...","SK_V5":"..."}},
  "work_focus":"...",
  "project_guidance":{{"project_key_1":["..."],"project_key_2":["..."]}}
}}
"""
    try:
        raw = _call_groq(system, user, GROQ_GEN_MODEL, max_tokens=1200)
        data = json.loads(_repair_json(raw))
    except Exception as exc:
        logger.warning("  AI tailoring strategy failed - using scored fallback: %s", exc)
        return fallback

    projects = []
    for key in data.get("projects", []):
        key = str(key).strip()
        if key in PROJECTS and key not in projects:
            projects.append(key)
    for key in fallback["projects"]:
        if key not in projects:
            projects.append(key)
        if len(projects) == 2:
            break
    projects = projects[:2]
    project_pair = (projects[0], projects[1])

    strategy = {
        "projects": projects,
        "skills": _normalize_skill_profile(data.get("skills", {}), fallback["skills"], project_pair),
        "work_focus": re.sub(r"\s+", " ", str(data.get("work_focus", fallback["work_focus"]))).strip()[:500],
        "project_guidance": {},
        "role_market": role_market_intel or {},
        "source": "ai",
    }
    raw_guidance = data.get("project_guidance", {}) if isinstance(data.get("project_guidance", {}), dict) else {}
    for key in projects:
        vals = raw_guidance.get(key, fallback.get("project_guidance", {}).get(key, []))
        if isinstance(vals, str):
            vals = [vals]
        strategy["project_guidance"][key] = [
            re.sub(r"\s+", " ", str(v)).strip()[:160]
            for v in vals[:3]
            if str(v).strip()
        ]
    logger.info("  AI tailoring strategy: projects=%s source=%s", " + ".join(projects), strategy["source"])
    return strategy


# ─────────────────────────────────────────────────────────────────────────────
# SOURCE-DRIVEN EXPERIENCE RESEARCH
# ─────────────────────────────────────────────────────────────────────────────
_HDRS = {"User-Agent": "Mozilla/5.0 (X11; Linux x86_64) Chrome/120.0.0.0 Safari/537.36"}

_RESEARCH_STOPWORDS = {
    "and","for","the","with","from","this","that","role","jobs","job",
    "resume","cv","india","bangalore","required","skills","experience",
    "analyst","engineer","associate","senior","junior","intern","trainee",
}

_DOMAIN_RESEARCH_TERMS = {
    "SOC":       ["soc","security operations","siem","alert triage","incident response"],
    "VAPT":      ["vapt","vulnerability","penetration testing","owasp","cvss"],
    "AppSec":    ["application security","appsec","owasp","sast","dast"],
    "GRC":       ["grc","governance risk compliance","iso 27001","nist","control testing"],
    "Risk":      ["technology risk","it risk","risk analyst","control testing","rcsa"],
    "Fraud-AML": ["fraud","aml","kyc","transaction monitoring","financial crime"],
    "CloudSec":  ["cloud security","aws security","iam","guardduty","cloudtrail"],
    "IAM":       ["iam","identity","access governance","pam","sailpoint","cyberark"],
    "Forensics": ["dfir","forensics","incident response","chain of custody"],
    "Network":   ["network security","firewall","ids","ips","packet analysis"],
    "General":   ["cybersecurity","information security","security analyst"],
}

_RESUME_ACTION_VERBS = {
    "analyzed","assessed","audited","automated","built","conducted","configured","created",
    "detected","developed","documented","enforced","evaluated","identified","implemented",
    "investigated","maintained","monitored","performed","prioritized","reviewed","scanned",
    "triaged","validated",
}

_PUBLIC_DOC_EXTS       = {".pdf",".docx",".doc"}
_MAX_PUBLIC_DOC_BYTES  = 3_000_000
_TECHNICAL_BENCHMARK_TERMS = {
    "splunk","python","bash","powershell","virustotal","abuseipdb","urlscan","wireshark",
    "nmap","aws","boto3","nessus","openvas","burp","metasploit","qradar","sentinel",
    "crowdstrike","defender","elastic","kibana","sysmon","guardduty","cloudtrail","cyberark",
    "sailpoint","okta","sql","linux","windows","dns","api","siem","osint",
    "threat intelligence","ioc","mitre","owasp",
}
_WORK_EXPERIENCE_SAFE_SIGNALS = (
    "triage","escalat","incident","investig","document","audit","evidence","root cause",
    "anomal","pattern","severity","risk","compliance","policy","control","fraud",
    "transaction","monitor","case","sla","workflow","handoff","priority","report",
)


def _scrub_external_text(text: str, max_chars: int = 12000) -> str:
    text = re.sub(r"\b[\w.+-]+@[\w.-]+\.\w+\b", " ", text or "")
    text = re.sub(r"\b(?:\+?\d[\d\s().-]{7,}\d)\b", " ", text)
    text = re.sub(r"https?://\S+|www\.\S+", " ", text)
    text = re.sub(r"\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}\b", " ", text, flags=re.I)
    text = re.sub(r"\s+", " ", text).strip()
    return text[:max_chars]


def _visible_html_text(html: str) -> str:
    soup = BeautifulSoup(html or "", "html.parser")
    for tag in soup(["script","style","nav","footer","header","form"]):
        tag.decompose()
    main = soup.find("main") or soup.find("article") or soup
    return " ".join(main.get_text(" ", strip=True).split())


def _extract_pdf_text_from_file(path: Path) -> str:
    try:
        from pdfminer.high_level import extract_text
        return extract_text(str(path)) or ""
    except Exception as exc:
        logger.debug("PDF text extraction failed for %s: %s", path.name, exc)
        return ""


def _extract_pdf_text_from_bytes(data: bytes) -> str:
    with tempfile.NamedTemporaryFile(suffix=".pdf") as tmp:
        tmp.write(data)
        tmp.flush()
        return _extract_pdf_text_from_file(Path(tmp.name))


def _extract_docx_text_from_bytes(data: bytes) -> str:
    try:
        doc = Document(io.BytesIO(data))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except Exception as exc:
        logger.debug("DOCX text extraction failed: %s", exc)
        return ""


def _role_research_terms(job: dict, jd_keywords: dict) -> list[str]:
    domain    = str(job.get("domain", "General")).strip() or "General"
    raw_terms = []
    raw_terms.extend(_DOMAIN_RESEARCH_TERMS.get(domain, _DOMAIN_RESEARCH_TERMS["General"]))
    raw_terms.extend(jd_keywords.get("ranked", [])[:10])
    raw_terms.extend(jd_keywords.get("tools", [])[:6])
    raw_terms.extend(re.findall(r"[a-zA-Z][a-zA-Z0-9+#.-]{2,}", str(job.get("job_title", ""))))
    raw_terms.extend(re.findall(r"[a-zA-Z][a-zA-Z0-9+#.-]{2,}", str(job.get("skills", "")))[:12])
    seen, terms = set(), []
    for term in raw_terms:
        term = re.sub(r"\s+", " ", str(term).lower()).strip(" -_/|")
        if len(term) < 3 or term in _RESEARCH_STOPWORDS or term in seen:
            continue
        seen.add(term)
        terms.append(term)
    return terms[:28]


def _source_relevance_score(text: str, terms: list[str]) -> int:
    lower = (text or "").lower()
    score = 0
    for term in terms:
        if not term:
            continue
        weight = 3 if " " in term else 1
        if re.search(rf"(?<!\w){re.escape(term)}(?!\w)", lower):
            score += weight
    return score


def _unwrap_duckduckgo_url(href: str) -> str:
    if not href:
        return ""
    if href.startswith("//"):
        href = "https:" + href
    parsed = urlparse(href)
    if "duckduckgo.com" in parsed.netloc and parsed.path.startswith("/l/"):
        uddg = parse_qs(parsed.query).get("uddg", [""])[0]
        if uddg:
            return unquote(uddg)
    return href


def _normalize_public_doc_url(url: str) -> str | None:
    if not url:
        return None
    parsed = urlparse(url.strip())
    if parsed.scheme not in {"http","https"} or not parsed.netloc:
        return None
    netloc = parsed.netloc.lower()
    path   = parsed.path
    if "linkedin.com" in netloc and not RESUME_RESEARCH_ALLOW_LINKEDIN:
        return None
    if netloc == "github.com" and "/blob/" in path:
        parts = path.strip("/").split("/")
        if len(parts) >= 5:
            owner, repo, _, branch = parts[:4]
            rest = "/".join(parts[4:])
            return f"https://raw.githubusercontent.com/{owner}/{repo}/{branch}/{rest}"
    suffix = Path(path).suffix.lower()
    if suffix in _PUBLIC_DOC_EXTS:
        return url
    if "raw.githubusercontent.com" in netloc and suffix in _PUBLIC_DOC_EXTS:
        return url
    return None


def _discover_public_resume_urls(job: dict, terms: list[str]) -> list[str]:
    if not WEB_RESUME_RESEARCH or RESUME_RESEARCH_MAX_WEB <= 0:
        return []
    query_terms = " ".join(terms[:6])
    query = (
        f'{job.get("job_title", "")} {query_terms} resume '
        "filetype:pdf OR filetype:docx github linkedin"
    )
    try:
        resp = requests.get(
            "https://html.duckduckgo.com/html/",
            params={"q": query[:220]},
            headers=_HDRS,
            timeout=8,
        )
        resp.raise_for_status()
    except Exception as exc:
        logger.debug("Public resume discovery failed: %s", exc)
        return []
    urls = []
    soup = BeautifulSoup(resp.text, "html.parser")
    for a in soup.select("a.result__a"):
        norm = _normalize_public_doc_url(_unwrap_duckduckgo_url(a.get("href", "")))
        if norm and norm not in urls:
            urls.append(norm)
        if len(urls) >= RESUME_RESEARCH_MAX_WEB:
            break
    return urls


def _fetch_public_document_text(url: str) -> str:
    try:
        resp = requests.get(url, headers=_HDRS, timeout=12, allow_redirects=True)
        resp.raise_for_status()
    except Exception as exc:
        logger.debug("Public resume fetch failed for %s: %s", url, exc)
        return ""
    if len(resp.content) > _MAX_PUBLIC_DOC_BYTES:
        logger.debug("Skipping oversized public resume: %s", url)
        return ""
    suffix = Path(urlparse(resp.url).path).suffix.lower()
    ctype  = resp.headers.get("content-type", "").lower()
    data   = resp.content
    if suffix == ".pdf" or "application/pdf" in ctype:
        text = _extract_pdf_text_from_bytes(data)
    elif suffix in {".docx",".doc"} or "wordprocessingml" in ctype:
        text = _extract_docx_text_from_bytes(data)
    elif "text/html" in ctype:
        text = _visible_html_text(resp.text)
    else:
        text = ""
    return _scrub_external_text(text)


def _load_public_resume_sources(job: dict, terms: list[str]) -> list[dict]:
    urls = []
    for url in RESUME_SOURCE_URLS:
        norm = _normalize_public_doc_url(url)
        if norm and norm not in urls:
            urls.append(norm)
    for url in _discover_public_resume_urls(job, terms):
        if url not in urls:
            urls.append(url)

    sources = []
    for url in urls[:RESUME_RESEARCH_MAX_WEB]:
        text = _fetch_public_document_text(url)
        if len(text) < 150:
            continue
        score = _source_relevance_score(text, terms)
        sources.append({
            "kind":   "public_resume",
            "source": urlparse(url).netloc,
            "text":   text,
            "score":  score,
        })

    # [CHANGE D] Filter low-relevance sources before passing to LLM prompt
    sources = [s for s in sources if s.get("score", 0) >= 3]
    sources.sort(key=lambda item: item.get("score", 0), reverse=True)
    return sources[:RESUME_RESEARCH_MAX_WEB]


def _fetch_reddit_forum_posts(job: dict, terms: list[str]) -> list[dict]:
    if not FORUM_RESEARCH or FORUM_RESEARCH_MAX_POSTS <= 0:
        return []
    subs    = REDDIT_RESEARCH_SUBS or ["cybersecurityindia","cybersecurity","AskNetsec"]
    per_sub = max(1, min(3, (FORUM_RESEARCH_MAX_POSTS + len(subs) - 1) // len(subs)))
    q_terms = " ".join([str(job.get("job_title","")), str(job.get("domain","")), *terms[:4]])
    query   = f"{q_terms} resume job career India".strip()
    posts   = []
    for sub in subs:
        if len(posts) >= FORUM_RESEARCH_MAX_POSTS:
            break
        try:
            resp = requests.get(
                f"https://www.reddit.com/r/{sub}/search.json",
                params={"q": query[:180], "restrict_sr":"1","sort":"relevance","limit":str(per_sub)},
                headers=_HDRS,
                timeout=8,
            )
            if resp.status_code in (403, 429):
                continue
            resp.raise_for_status()
            data = resp.json()
        except Exception as exc:
            logger.debug("Reddit research failed for r/%s: %s", sub, exc)
            continue
        for child in data.get("data", {}).get("children", []):
            post  = child.get("data", {})
            title = post.get("title", "")
            body  = post.get("selftext", "")
            text  = _scrub_external_text(f"{title}. {body}", max_chars=1200)
            if len(text) < 40:
                continue
            posts.append({
                "kind":   "forum",
                "source": f"r/{sub}",
                "text":   text,
                "score":  _source_relevance_score(text, terms),
            })
            if len(posts) >= FORUM_RESEARCH_MAX_POSTS:
                break
    posts.sort(key=lambda item: item.get("score", 0), reverse=True)
    return posts[:FORUM_RESEARCH_MAX_POSTS]


def _extract_resume_bullet_lines(texts: list[str], terms: list[str]) -> list[str]:
    bullets  = []
    term_blob = "|".join(re.escape(t) for t in terms[:18] if len(t) > 2)
    term_re   = re.compile(term_blob, re.IGNORECASE) if term_blob else None
    for text in texts:
        for raw in re.split(r"[\n\r]+|(?<=\.)\s+(?=[A-Z][a-z]+(?:ed|d|ing)\b)", text):
            line = raw.strip(" \t-*•·")
            if not (45 <= len(line) <= 280):
                continue
            first = re.match(r"([A-Za-z]+)", line.lower())
            starts_action = bool(first and first.group(1) in _RESUME_ACTION_VERBS)
            has_term      = bool(term_re and term_re.search(line))
            if starts_action or has_term:
                bullets.append(line)
            if len(bullets) >= 80:
                return bullets
    return bullets


def _top_signal_terms(texts: list[str], seed_terms: list[str], limit: int = 10) -> list[str]:
    counts = Counter()
    for text in texts:
        lower = text.lower()
        for term in seed_terms:
            if len(term) < 3:
                continue
            count = len(re.findall(rf"(?<!\w){re.escape(term)}(?!\w)", lower, flags=re.IGNORECASE))
            if count:
                counts[term] += count
    return [term for term, _ in counts.most_common(limit)]


def _top_action_verbs(bullets: list[str], limit: int = 7) -> list[str]:
    counts = Counter()
    for bullet in bullets:
        m = re.match(r"([A-Za-z]+)", bullet.lower())
        if m and m.group(1) in _RESUME_ACTION_VERBS:
            counts[m.group(1)] += 1
    return [verb for verb, _ in counts.most_common(limit)]


def _metric_style_signals(texts: list[str]) -> list[str]:
    blob    = " ".join(texts).lower()
    signals = []
    if re.search(r"\b\d+\+?\s*(?:alerts|incidents|tickets|cases|events|logs|vulnerabilities)\b", blob):
        signals.append("quantified queue/alert/case volume")
    if re.search(r"\b(?:sla|mttr|tat|deadline|within\s+\d+\s*(?:hours|days|hrs))\b", blob):
        signals.append("SLA or turnaround-time language")
    if re.search(r"\b(?:audit|evidence|control|compliance|documentation|reporting)\b", blob):
        signals.append("audit evidence and documentation outcomes")
    if re.search(r"\b(?:false positive|triage|escalation|priority|severity)\b", blob):
        signals.append("severity-based triage and escalation")
    if re.search(r"\b(?:reduced|improved|automated|decreased|increased|optimized)\b", blob):
        signals.append("result-led impact wording")
    return signals[:5]


def _build_transition_bridges(target_terms: list[str]) -> list[str]:
    target_blob = " ".join(target_terms).lower()
    bridges     = []
    if re.search(r"\b(alert|triage|incident|soc|monitor|ticket|severity)\b", target_blob):
        bridges.append("case triage -> alert triage, severity classification, and escalation routing")
    if re.search(r"\b(audit|compliance|control|evidence|grc|risk|sox|iso|nist)\b", target_blob):
        bridges.append("audit-ready case notes -> control evidence, compliance tracking, and decision trails")
    if re.search(r"\b(fraud|aml|kyc|transaction|suspicious|anomaly|threat|indicator)\b", target_blob):
        bridges.append("seller-claim anomalies -> suspicious pattern detection and investigation handoffs")
    if re.search(r"\b(vulnerability|patch|remediation|sla|cvss|risk)\b", target_blob):
        bridges.append("policy exception handling -> risk prioritization, SLA discipline, and remediation tracking")
    if re.search(r"\b(iam|access|identity|cloud|policy|governance)\b", target_blob):
        bridges.append("policy eligibility review -> access governance mindset and exception documentation")
    if not bridges:
        bridges.append("high-volume operations review -> structured investigation, escalation, and evidence documentation")
    return bridges[:4]


def _is_work_experience_safe_term(term: str) -> bool:
    lower = term.lower()
    if lower in _TECHNICAL_BENCHMARK_TERMS:
        return False
    return any(signal in lower for signal in _WORK_EXPERIENCE_SAFE_SIGNALS)


# ─────────────────────────────────────────────────────────────────────────────
# ROLE-MARKET INTELLIGENCE  (Reddit + GitHub + web snippets)
# ─────────────────────────────────────────────────────────────────────────────
def _github_headers() -> dict:
    headers = {"Accept": "application/vnd.github+json"}
    if GITHUB_TOKEN:
        headers["Authorization"] = f"Bearer {GITHUB_TOKEN}"
    return headers


def _clean_market_list(values, limit: int = 10, max_chars: int = 80) -> list[str]:
    if isinstance(values, str):
        values = re.split(r"[,;\n]+", values)
    if not isinstance(values, list):
        return []
    cleaned, seen = [], set()
    for raw in values:
        item = re.sub(r"\s+", " ", str(raw)).strip(" .;:-")
        if not item or len(item) > max_chars:
            continue
        key = item.lower()
        if key in seen:
            continue
        seen.add(key)
        cleaned.append(item)
        if len(cleaned) >= limit:
            break
    return cleaned


def _market_query_terms(job: dict, jd_keywords: dict) -> list[str]:
    terms = _role_research_terms(job, jd_keywords)
    role_terms = re.findall(r"[a-zA-Z][a-zA-Z0-9+#.-]{2,}", str(job.get("job_title", "")))
    raw = role_terms + terms + jd_keywords.get("ranked", [])[:10] + jd_keywords.get("tools", [])[:6]
    cleaned, seen = [], set()
    for term in raw:
        term = re.sub(r"\s+", " ", str(term).lower()).strip(" -_/|")
        if len(term) < 3 or term in _RESEARCH_STOPWORDS or term in seen:
            continue
        seen.add(term)
        cleaned.append(term)
    return cleaned[:24]


def _fetch_market_web_sources(job: dict, terms: list[str]) -> list[dict]:
    if not WEB_RESUME_RESEARCH or ROLE_MARKET_MAX_WEB <= 0:
        return []
    title = str(job.get("job_title", "")).strip()
    domain = str(job.get("domain", "")).strip()
    queries = [
        f'{title} {domain} entry level skills projects resume cybersecurity',
        f'{title} {domain} reddit skills projects fresher India',
    ]
    sources = []
    for query in queries:
        if len(sources) >= ROLE_MARKET_MAX_WEB:
            break
        try:
            resp = requests.get(
                "https://html.duckduckgo.com/html/",
                params={"q": query[:220]},
                headers=_HDRS,
                timeout=8,
            )
            resp.raise_for_status()
        except Exception as exc:
            logger.debug("Role-market web search failed: %s", exc)
            continue
        soup = BeautifulSoup(resp.text, "html.parser")
        for result in soup.select(".result"):
            link = result.select_one("a.result__a")
            snippet = result.select_one(".result__snippet")
            if not link:
                continue
            url = _unwrap_duckduckgo_url(link.get("href", ""))
            host = urlparse(url).netloc.lower()
            if any(blocked in host for blocked in ("linkedin.com", "indeed.com", "glassdoor.com")):
                continue
            text = _scrub_external_text(
                f"{link.get_text(' ', strip=True)}. {snippet.get_text(' ', strip=True) if snippet else ''}",
                max_chars=900,
            )
            if len(text) < 50:
                continue
            sources.append({
                "kind": "market_web",
                "source": host or "duckduckgo",
                "text": text,
                "score": _source_relevance_score(text, terms),
            })
            if len(sources) >= ROLE_MARKET_MAX_WEB:
                break
    sources.sort(key=lambda item: item.get("score", 0), reverse=True)
    return sources[:ROLE_MARKET_MAX_WEB]


def _fetch_github_readme_snippet(full_name: str) -> str:
    try:
        resp = requests.get(
            f"https://api.github.com/repos/{full_name}/readme",
            headers=_github_headers(),
            timeout=8,
        )
        if resp.status_code in (403, 404, 429):
            return ""
        resp.raise_for_status()
        payload = resp.json()
        encoded = payload.get("content", "")
        if not encoded:
            return ""
        raw = base64.b64decode(encoded).decode("utf-8", errors="ignore")
        return _scrub_markdown_text(raw, max_chars=1400)
    except Exception as exc:
        logger.debug("GitHub README market fetch failed for %s: %s", full_name, exc)
        return ""


def _fetch_market_github_sources(job: dict, terms: list[str]) -> list[dict]:
    if not PROJECT_GITHUB_RESEARCH or ROLE_MARKET_MAX_GITHUB <= 0:
        return []
    title = re.sub(r"[^a-zA-Z0-9 +#.-]+", " ", str(job.get("job_title", ""))).strip()
    domain_terms = " ".join(terms[:6])
    query = f"{title} {domain_terms} cybersecurity portfolio project in:readme stars:>1"
    try:
        resp = requests.get(
            "https://api.github.com/search/repositories",
            params={"q": query[:240], "sort": "stars", "per_page": str(ROLE_MARKET_MAX_GITHUB)},
            headers=_github_headers(),
            timeout=10,
        )
        if resp.status_code in (403, 429):
            logger.debug("GitHub market search rate-limited: %s", resp.status_code)
            return []
        resp.raise_for_status()
        items = resp.json().get("items", [])
    except Exception as exc:
        logger.debug("GitHub market search failed: %s", exc)
        return []

    sources = []
    for item in items[:ROLE_MARKET_MAX_GITHUB]:
        full_name = item.get("full_name", "")
        readme = _fetch_github_readme_snippet(full_name) if full_name else ""
        text = _scrub_external_text(
            " ".join(filter(None, [
                full_name,
                item.get("description") or "",
                " ".join(item.get("topics", [])[:10]),
                readme,
            ])),
            max_chars=2200,
        )
        if len(text) < 60:
            continue
        sources.append({
            "kind": "market_github",
            "source": full_name,
            "text": text,
            "score": _source_relevance_score(text, terms) + min(int(item.get("stargazers_count", 0)), 50) // 10,
        })
    sources.sort(key=lambda item: item.get("score", 0), reverse=True)
    return sources[:ROLE_MARKET_MAX_GITHUB]


def _fetch_market_reddit_sources(job: dict, terms: list[str]) -> list[dict]:
    if not FORUM_RESEARCH or ROLE_MARKET_MAX_REDDIT <= 0:
        return []
    subs = REDDIT_RESEARCH_SUBS or ["cybersecurityindia", "cybersecurity", "AskNetsec"]
    per_sub = max(1, min(3, (ROLE_MARKET_MAX_REDDIT + len(subs) - 1) // len(subs)))
    title = str(job.get("job_title", ""))
    query = f"{title} skills projects resume entry level fresher India cybersecurity".strip()
    sources = []
    for sub in subs:
        if len(sources) >= ROLE_MARKET_MAX_REDDIT:
            break
        try:
            resp = requests.get(
                f"https://www.reddit.com/r/{sub}/search.json",
                params={"q": query[:180], "restrict_sr": "1", "sort": "relevance", "limit": str(per_sub)},
                headers=_HDRS,
                timeout=8,
            )
            if resp.status_code in (403, 429):
                continue
            resp.raise_for_status()
            data = resp.json()
        except Exception as exc:
            logger.debug("Role-market Reddit search failed for r/%s: %s", sub, exc)
            continue
        for child in data.get("data", {}).get("children", []):
            post = child.get("data", {})
            text = _scrub_external_text(
                f"{post.get('title', '')}. {post.get('selftext', '')}",
                max_chars=1400,
            )
            if len(text) < 60:
                continue
            sources.append({
                "kind": "market_reddit",
                "source": f"r/{sub}",
                "text": text,
                "score": _source_relevance_score(text, terms),
            })
            if len(sources) >= ROLE_MARKET_MAX_REDDIT:
                break
    sources.sort(key=lambda item: item.get("score", 0), reverse=True)
    return sources[:ROLE_MARKET_MAX_REDDIT]


def _project_angles_from_terms(terms: list[str]) -> list[str]:
    blob = " ".join(terms).lower()
    patterns = [
        (r"\b(soc|siem|splunk|alert|incident|detection|sigma|soar)\b", "SIEM alert triage, detection rules, and incident escalation"),
        (r"\b(vulnerability|cve|cvss|epss|nessus|openvas|patch|remediation)\b", "vulnerability prioritization, CVE evidence, and remediation tracking"),
        (r"\b(appsec|application security|owasp|sast|dast|burp|secure code)\b", "OWASP testing, application security evidence, and secure remediation notes"),
        (r"\b(grc|audit|compliance|control|nist|iso|sox|itgc)\b", "audit-ready evidence, control mapping, and compliance documentation"),
        (r"\b(cloud|aws|iam|guardduty|cloudtrail|cspm|identity)\b", "cloud/IAM monitoring, access-risk review, and policy exception evidence"),
        (r"\b(osint|phishing|ioc|threat intelligence|virustotal|urlscan|abuseipdb)\b", "OSINT enrichment, IOC investigation, and phishing infrastructure analysis"),
        (r"\b(fraud|aml|kyc|transaction|sanctions|financial crime)\b", "fraud-pattern review, KYC evidence, and suspicious-activity investigation workflows"),
        (r"\b(network|firewall|ids|ips|packet|wireshark|tcp)\b", "network traffic analysis, packet evidence, and IDS alert correlation"),
    ]
    angles = []
    for pattern, angle in patterns:
        if re.search(pattern, blob) and angle not in angles:
            angles.append(angle)
    return angles[:6]


def _market_avoid_claims(market_terms: list[str], texts: list[str]) -> list[str]:
    blob = " ".join(texts + market_terms).lower()
    avoid = []
    for term in sorted(_CANDIDATE_CANNOT_MEET):
        if re.search(rf"(?<!\w){re.escape(term)}(?!\w)", blob):
            avoid.append(term)
    evidence_blob = _candidate_evidence_blob(tuple(PROJECTS.keys()))
    for term in market_terms:
        lower = term.lower()
        technical = lower in _TECHNICAL_BENCHMARK_TERMS or any(
            signal in lower for signal in ("certified", "certification", "cissp", "cisa", "oscp", "giac", "senior", "lead")
        )
        if technical and not _skill_item_grounded(term, evidence_blob):
            avoid.append(term)
    return _clean_market_list(avoid, limit=10)


def _role_market_llm_summary(job: dict, jd_keywords: dict, sources: list[dict], fallback: dict) -> dict:
    if not ROLE_MARKET_LLM_SUMMARY or not AI_TAILORING or not GROQ_API_KEY or not sources:
        return fallback
    snippets = "\n\n".join(
        f"[{s.get('kind')}:{s.get('source')}] {s.get('text', '')[:900]}"
        for s in sources[:10]
    )
    system = (
        "You are a cybersecurity role-market analyst. Return ONLY valid JSON. "
        "Separate market demand from what the candidate can honestly claim."
    )
    user = f"""
Target job: {job.get('job_title', '')} | Domain: {job.get('domain', '')}
JD summary/skills: {job.get('summary', '')} {job.get('skills', '')}
JD extracted keywords: {', '.join(jd_keywords.get('ranked', [])[:12])}

Candidate evidence summary:
- Current role: {CURRENT_ROLE_TITLE}
- Current facts: {'; '.join(AMAZON_RAW_FACTS)}
- Candidate/project vocabulary: {_CANDIDATE_PROFILE[:900]}
- Project titles: {'; '.join(p['title'] for p in PROJECTS.values())}

Reddit/GitHub/web role-market snippets:
{snippets[:6000]}

Return JSON:
{{
  "target_skills":["role-market skill/tool/concept", "..."],
  "experience_keywords":["transferable work-experience keyword", "..."],
  "project_angles":["project storyline recruiters expect", "..."],
  "project_keywords":["project/tool/concept keyword", "..."],
  "avoid_claiming":["unsupported senior/cert/tool claim", "..."],
  "role_summary":"one concise sentence"
}}

Rules:
- target_skills are what the market/JD wants, but they still must be filtered before the final resume.
- experience_keywords must be safe for operations-to-security transfer: triage, investigation, evidence, policy, audit, escalation, patterns, risk.
- project_angles should describe what project bullets should emphasize.
- avoid_claiming should include seniority, certifications, platforms, or tools not grounded in candidate/project evidence.
"""
    try:
        raw = _call_groq(system, user, GROQ_GEN_MODEL, max_tokens=900)
        data = json.loads(_repair_json(raw))
    except Exception as exc:
        logger.debug("Role-market LLM summary failed: %s", exc)
        return fallback

    result = dict(fallback)
    result["target_skills"] = _clean_market_list(data.get("target_skills", []), 14) or fallback.get("target_skills", [])
    result["experience_keywords"] = _clean_market_list(data.get("experience_keywords", []), 12) or fallback.get("experience_keywords", [])
    result["project_angles"] = _clean_market_list(data.get("project_angles", []), 8, 140) or fallback.get("project_angles", [])
    result["project_keywords"] = _clean_market_list(data.get("project_keywords", []), 12) or fallback.get("project_keywords", [])
    avoid = _clean_market_list(data.get("avoid_claiming", []), 10)
    result["avoid_claiming"] = _clean_market_list(avoid + fallback.get("avoid_claiming", []), 10)
    role_summary = re.sub(r"\s+", " ", str(data.get("role_summary", ""))).strip()
    if role_summary:
        result["role_summary"] = role_summary[:300]
    result["summary"] = _format_role_market_summary(result)
    return result


def _format_role_market_summary(intel: dict) -> str:
    parts = []
    if intel.get("role_summary"):
        parts.append(intel["role_summary"])
    if intel.get("target_skills"):
        parts.append("Target market skills: " + ", ".join(intel["target_skills"][:12]) + ".")
    if intel.get("experience_keywords"):
        parts.append("Experience keywords to weave: " + ", ".join(intel["experience_keywords"][:10]) + ".")
    if intel.get("project_angles"):
        parts.append("Project angles: " + "; ".join(intel["project_angles"][:6]) + ".")
    if intel.get("avoid_claiming"):
        parts.append("Avoid unsupported claims: " + ", ".join(intel["avoid_claiming"][:8]) + ".")
    if intel.get("source_note"):
        parts.append("Sources: " + intel["source_note"] + ".")
    return " ".join(parts)


def _summarize_role_market_intel(job: dict, jd_keywords: dict, sources: list[dict]) -> dict:
    texts = [s.get("text", "") for s in sources if s.get("text")]
    terms = _market_query_terms(job, jd_keywords)
    seed_terms = list(dict.fromkeys(
        terms
        + jd_keywords.get("ranked", [])[:14]
        + jd_keywords.get("tools", [])[:8]
        + list(CANDIDATE_GROUNDABLE)
        + list(_TECHNICAL_BENCHMARK_TERMS)
    ))
    market_terms = _top_signal_terms(texts + [_jd_blob(job, jd_keywords)], seed_terms, limit=22)
    experience_keywords = [t for t in market_terms if _is_work_experience_safe_term(t)]
    project_keywords = [t for t in market_terms if t not in experience_keywords]
    project_angles = _project_angles_from_terms(market_terms + jd_keywords.get("ranked", []))
    source_counts = Counter(s.get("kind", "unknown") for s in sources)
    source_note = (
        f"web={source_counts.get('market_web', 0)}, "
        f"github={source_counts.get('market_github', 0)}, "
        f"reddit={source_counts.get('market_reddit', 0)}"
    )
    fallback = {
        "role_summary": f"Market scan for {job.get('job_title', 'target role')} prioritized role keywords, fresher-safe project angles, and transferable experience wording.",
        "target_skills": market_terms[:14],
        "experience_keywords": experience_keywords[:12] or _build_transition_bridges(market_terms + jd_keywords.get("ranked", [])),
        "project_angles": project_angles,
        "project_keywords": project_keywords[:12],
        "avoid_claiming": _market_avoid_claims(market_terms, texts),
        "source_note": source_note,
        "sources": [
            {"kind": s.get("kind"), "source": s.get("source"), "score": s.get("score", 0)}
            for s in sources[:8]
        ],
    }
    fallback["summary"] = _format_role_market_summary(fallback)
    return _role_market_llm_summary(job, jd_keywords, sources, fallback)


def collect_role_market_intel(job: dict, jd_keywords: dict) -> dict:
    if not ROLE_MARKET_RESEARCH:
        return {"summary": "", "source_note": "disabled"}
    cache_key = f"{job.get('domain', 'General')}:{job.get('job_title', '')[:32]}:{','.join(jd_keywords.get('ranked', [])[:3])}"
    if cache_key in _ROLE_MARKET_CACHE:
        logger.info("  Role market intel: cache hit (%s)", cache_key[:60])
        return _ROLE_MARKET_CACHE[cache_key]

    terms = _market_query_terms(job, jd_keywords)
    sources = []
    sources.extend(_fetch_market_web_sources(job, terms))
    sources.extend(_fetch_market_github_sources(job, terms))
    sources.extend(_fetch_market_reddit_sources(job, terms))
    sources = [s for s in sources if s.get("score", 0) >= 1 or s.get("kind") == "market_github"]
    sources.sort(key=lambda item: item.get("score", 0), reverse=True)
    intel = _summarize_role_market_intel(job, jd_keywords, sources)
    logger.info("  Role market intel: %s", intel.get("source_note", ""))
    _ROLE_MARKET_CACHE[cache_key] = intel
    return intel


def _summarize_experience_research(
    job: dict, jd_keywords: dict,
    resume_sources: list[dict], forum_sources: list[dict]
) -> dict:
    resume_texts = [s["text"] for s in resume_sources if s.get("text")]
    forum_texts  = [s["text"] for s in forum_sources  if s.get("text")]
    terms        = _role_research_terms(job, jd_keywords)
    seed_terms   = list(dict.fromkeys(
        terms
        + list(CANDIDATE_GROUNDABLE)
        + list(_CANDIDATE_CAN_MEET)
        + jd_keywords.get("ranked", [])[:12]
    ))

    bullets          = _extract_resume_bullet_lines(resume_texts, terms)
    target_terms     = _top_signal_terms(resume_texts, seed_terms, limit=12)
    experience_terms = [t for t in target_terms if _is_work_experience_safe_term(t)]
    technical_terms  = [t for t in target_terms if t not in experience_terms]
    forum_terms      = _top_signal_terms(forum_texts, seed_terms, limit=8)
    verbs            = _top_action_verbs(bullets, limit=7)
    metric_styles    = _metric_style_signals(resume_texts + forum_texts)
    bridges          = _build_transition_bridges(target_terms + forum_terms + jd_keywords.get("ranked", []))

    source_counts = Counter(s["kind"] for s in resume_sources + forum_sources)
    source_note   = (
        "local resumes=ignored, "
        f"public resumes={source_counts.get('public_resume', 0)}, "
        f"forums={source_counts.get('forum', 0)}"
    )

    summary_parts = [
        f"Target role: {job.get('job_title', 'Role')} ({job.get('domain', 'General')}).",
        f"Transition from: {CURRENT_ROLE_TITLE}.",
    ]
    if target_terms:
        summary_parts.append("Benchmark resume language: " + ", ".join(target_terms[:10]) + ".")
    if experience_terms:
        summary_parts.append("Experience-safe framing terms: " + ", ".join(experience_terms[:8]) + ".")
    if technical_terms:
        summary_parts.append("Technical/project terms: " + ", ".join(technical_terms[:8]) + ".")
    if forum_terms:
        summary_parts.append("Forum/reddit market signals: " + ", ".join(forum_terms[:7]) + ".")
    if verbs:
        summary_parts.append("Common action verbs: " + ", ".join(verbs[:7]) + ".")
    if metric_styles:
        summary_parts.append("Bullet style signals: " + ", ".join(metric_styles[:5]) + ".")
    if bridges:
        summary_parts.append("Transition bridges: " + "; ".join(bridges) + ".")
    summary_parts.append(
        "Use these as framing signals only; do not copy wording or borrow tools/certs from external resumes."
    )

    return {
        "summary":          " ".join(summary_parts),
        "target_terms":     target_terms,
        "experience_terms": experience_terms,
        "technical_terms":  technical_terms,
        "forum_terms":      forum_terms,
        "action_verbs":     verbs,
        "bridges":          bridges,
        "source_note":      source_note,
    }


def collect_experience_research(job: dict, jd_keywords: dict) -> dict:
    """
    Collect public resume and forum signals for the given job.
    [CHANGE E] Results are cached in _RESEARCH_CACHE keyed by domain:title[:20].
    Subsequent jobs with the same domain+title pattern skip HTTP calls entirely.
    """
    if not PUBLIC_RESUME_RESEARCH:
        return {"summary": "", "source_note": "disabled"}

    # [CHANGE E] Cache check
    cache_key = f"{job.get('domain', 'General')}:{job.get('job_title', '')[:20]}"
    if cache_key in _RESEARCH_CACHE:
        logger.info("  Experience research: cache hit (%s)", cache_key)
        return _RESEARCH_CACHE[cache_key]

    terms          = _role_research_terms(job, jd_keywords)
    public_sources = _load_public_resume_sources(job, terms)
    forum_sources  = _fetch_reddit_forum_posts(job, terms)
    research       = _summarize_experience_research(job, jd_keywords, public_sources, forum_sources)
    logger.info("  Experience research: %s", research.get("source_note", ""))

    # [CHANGE E] Store in cache
    _RESEARCH_CACHE[cache_key] = research
    return research


def format_experience_research_prompt(experience_research: dict | None,
                                      for_work_experience: bool = False) -> str:
    if not experience_research or not experience_research.get("summary"):
        return ""
    if for_work_experience:
        safe_terms   = ", ".join(experience_research.get("experience_terms", [])[:8]) or "triage, investigation, documentation, escalation"
        forum_terms  = ", ".join(experience_research.get("forum_terms", [])[:6])
        action_verbs = ", ".join(experience_research.get("action_verbs", [])[:6])
        bridges      = "; ".join(experience_research.get("bridges", [])[:4])
        return f"""
SOURCE-DRIVEN WORK-EXPERIENCE BENCHMARKS:
Target role: use transferable framing only for {CURRENT_ROLE_TITLE}.
Experience-safe framing terms: {safe_terms}.
Forum/reddit market signals: {forum_terms}.
Common action verbs: {action_verbs}.
Transition bridges: {bridges}.

BENCHMARK USAGE RULES:
- Use this to compare current-role facts against the target role, then choose the closest transferable framing.
- Do not copy external resume wording.
- Do not add tools, certifications, platforms, or achievements from benchmark resumes to the work-experience bullets.
- For work experience, only use the raw current-role facts and transferable language.
"""
    return f"""
SOURCE-DRIVEN ROLE BENCHMARKS:
{experience_research['summary'][:1800]}

BENCHMARK USAGE RULES:
- Compare the target role against the current role facts, then choose the closest transferable framing.
- Do not copy external resume wording.
- Do not add tools, certifications, platforms, or achievements from benchmark resumes to the work-experience bullets.
- For work experience, only use the raw current-role facts and transferable language.
"""


def scrape_company(company_raw: str) -> str:
    name = re.sub(r"\s*\(.*?\)\s*$", "", company_raw).strip()
    if not name or name.lower() in ("unknown", ""):
        return ""
    try:
        q    = requests.utils.quote(f"{name} cybersecurity about mission")
        resp = requests.get(f"https://html.duckduckgo.com/html/?q={q}", headers=_HDRS, timeout=8)
        soup = BeautifulSoup(resp.text, "html.parser")
        for a in soup.select("a.result__a"):
            href = a.get("href", "")
            if href.startswith("http") and not any(x in href for x in ["linkedin.com","glassdoor.com","indeed.com"]):
                pg   = requests.get(href, headers=_HDRS, timeout=8)
                s2   = BeautifulSoup(pg.text, "html.parser")
                for tag in s2(["script","style","nav","footer","header"]):
                    tag.decompose()
                main = s2.find("main") or s2.find("article") or s2
                text = " ".join(p.get_text(" ", strip=True) for p in main.find_all("p") if len(p.get_text()) > 40)
                if len(text) > 100:
                    return text[:800]
    except Exception:
        pass
    return ""


def research_github_projects(domain: str, job_title: str) -> str:
    DOMAIN_SEARCH = {
        "SOC":"SOC automation SIEM detection lab","VAPT":"vulnerability scanner CVE CVSS python",
        "GRC":"GRC compliance automation NIST ISO27001 python","Risk":"risk management compliance python",
        "Fraud-AML":"AML transaction monitoring fraud detection python",
        "CloudSec":"cloud security AWS IAM audit python","General":"cybersecurity portfolio entry level",
    }
    query   = DOMAIN_SEARCH.get(domain, "cybersecurity portfolio")
    encoded = requests.utils.quote(f"{query} language:Python stars:>2")
    url     = f"https://api.github.com/search/repositories?q={encoded}&sort=stars&per_page=5"
    headers = {"Accept":"application/vnd.github+json"}
    if GITHUB_TOKEN:
        headers["Authorization"] = f"Bearer {GITHUB_TOKEN}"
    try:
        resp  = requests.get(url, headers=headers, timeout=10)
        resp.raise_for_status()
        items = resp.json().get("items", [])
        return "\n".join(
            f"{i.get('full_name','')} (⭐{i.get('stargazers_count',0)}): "
            f"{(i.get('description','') or '')[:80]} | topics: {', '.join(i.get('topics',[])[:5])}"
            for i in items[:4]
        )
    except Exception as exc:
        logger.debug("GitHub research failed: %s", exc)
        return ""


# ─────────────────────────────────────────────────────────────────────────────
# JSON REPAIR + GROQ
# ─────────────────────────────────────────────────────────────────────────────
def _repair_json(raw: str) -> str:
    raw = re.sub(r"^```(?:json)?\s*", "", raw.strip())
    raw = re.sub(r"\s*```$", "", raw.strip())
    raw = raw.replace("\u201c", '"').replace("\u201d", '"')
    raw = raw.replace("\u2018", "'").replace("\u2019", "'")
    raw = re.sub(r",\s*([\}\]])", r"\1", raw)
    raw = re.sub(r'\\([^"\\/bfnrtu])', r'\1', raw)
    return raw.strip()


def _call_groq(system: str, user: str, model: str,
               max_tokens: int = 2500, retries: int | None = None) -> str:
    global _groq_call_count, _groq_consecutive_429s, _groq_cooldown_until, _groq_last_call_ts
    retries = GROQ_MAX_RETRIES if retries is None else max(1, retries)
    if not GROQ_API_KEY:
        raise RuntimeError("GROQ_API_KEY not set.")

    now = time.time()
    if _groq_cooldown_until and now < _groq_cooldown_until:
        remaining = int(_groq_cooldown_until - now)
        raise RuntimeError(f"Groq cooldown active for {remaining}s after repeated 429s.")

    payload = {
        "model": model, "temperature": 0.15, "max_tokens": max_tokens,
        "messages": [{"role":"system","content":system},{"role":"user","content":user}],
    }
    hdrs = {"Authorization": f"Bearer {GROQ_API_KEY}", "Content-Type": "application/json"}
    for attempt in range(1, retries + 1):
        try:
            with _GROQ_LOCK:
                now = time.time()
                if _groq_cooldown_until and now < _groq_cooldown_until:
                    remaining = int(_groq_cooldown_until - now)
                    raise RuntimeError(f"Groq cooldown active for {remaining}s after repeated 429s.")
                elapsed = now - _groq_last_call_ts
                if GROQ_MIN_INTERVAL_SECONDS > 0 and elapsed < GROQ_MIN_INTERVAL_SECONDS:
                    time.sleep(GROQ_MIN_INTERVAL_SECONDS - elapsed)
                _groq_last_call_ts = time.time()
                r = requests.post(GROQ_URL, json=payload, headers=hdrs, timeout=35)
            if r.status_code == 429:
                _groq_consecutive_429s += 1
                if GROQ_COOLDOWN_AFTER_429S and _groq_consecutive_429s >= GROQ_COOLDOWN_AFTER_429S:
                    _groq_cooldown_until = time.time() + GROQ_COOLDOWN_SECONDS
                    raise RuntimeError(
                        f"Groq rate-limit circuit opened after {_groq_consecutive_429s} consecutive 429s."
                    )
                wait = GROQ_429_WAIT_BASE * attempt
                logger.warning(
                    "  Groq 429 — %s (attempt %d/%d, model=%s)",
                    f"waiting {wait:.0f}s" if attempt < retries and wait > 0 else "no retry wait",
                    attempt, retries, model,
                )
                if attempt < retries and wait > 0:
                    time.sleep(wait)
                continue
            r.raise_for_status()
            _groq_consecutive_429s = 0
            _groq_call_count += 1
            return r.json()["choices"][0]["message"]["content"].strip()
        except requests.RequestException as exc:
            logger.warning("  Groq error attempt %d/%d (%s): %s", attempt, retries, model, exc)
            if attempt < retries:
                time.sleep(5 * attempt)
    raise RuntimeError(f"Groq ({model}) failed after retries.")


# ─────────────────────────────────────────────────────────────────────────────
# RESUME CONTENT GENERATION
# ─────────────────────────────────────────────────────────────────────────────
def generate_content(job: dict, p1_key: str, p2_key: str,
                     intel: dict | None, scraped_ctx: str,
                     p1_tools: list, p2_tools: list,
                     jd_keywords: dict,
                     experience_research: dict | None = None,
                     tailoring_strategy: dict | None = None,
                     project_evidence: dict | None = None) -> dict:

    # Phase A: Amazon bullets (separate Groq call, outcome-first)
    amazon_bullets = generate_amazon_bullets_dynamic(
        job, jd_keywords, experience_research, tailoring_strategy
    )

    # Phase B: Gap analysis + lens context for project generation
    jd_text_for_gap = f"{job.get('skills', '')} {job.get('summary', '')}"
    gap         = jd_gap_analysis(jd_text_for_gap, jd_keywords)
    lens_ctx    = build_three_lens_context(job.get("domain", "General"), jd_text_for_gap)
    research_ctx = format_experience_research_prompt(experience_research)

    p1 = PROJECTS[p1_key]
    p2 = PROJECTS[p2_key]
    project_evidence = project_evidence or {
        p1_key: get_project_evidence(p1_key, job.get("domain", "General")),
        p2_key: get_project_evidence(p2_key, job.get("domain", "General")),
    }
    guidance = (tailoring_strategy or {}).get("project_guidance", {})
    p1_guidance = "; ".join(guidance.get(p1_key, []) or select_concepts(p1_key, jd_text_for_gap))
    p2_guidance = "; ".join(guidance.get(p2_key, []) or select_concepts(p2_key, jd_text_for_gap))
    role_market = (tailoring_strategy or {}).get("role_market", {}) or {}
    market_project_angles = "; ".join(role_market.get("project_angles", [])[:8])
    market_project_keywords = ", ".join(role_market.get("project_keywords", [])[:12])

    co_ctx = ""
    if intel:
        co_ctx = f"\nCOMPANY FRAMING: {intel['framing']}\nPriority keywords: {', '.join(intel['keywords'][:4])}\nDo NOT write 'Eager to contribute to X'.\n"
    elif scraped_ctx:
        co_ctx = f"\nCOMPANY CONTEXT: {scraped_ctx[:400]}\n"

    ranked = jd_keywords.get("ranked", [])
    kw_hint = ""
    if ranked:
        kw_hint = (
            f"\nKEYWORD INJECTION: Weave these top JD keywords naturally across bullets "
            f"(target 2-3x total, max 2 per bullet): {', '.join(ranked[:8])}\n"
        )

    system = (
        "You are a senior cybersecurity resume writer for the Indian job market. "
        "Bullets must be factual — never fabricate tools or experience. "
        "ALWAYS write 'and' not '&' in bullet text (except MITRE ATT&CK which is a proper noun). "
        "BULLET END RULE: every bullet must end with a concrete technical artifact, workflow, or finding. "
        "NEVER end with 'to improve X', 'to optimize Y', 'to ensure Z', "
        "'to streamline X', 'to strengthen X', or 'in order to X'. "
        "Return ONLY a valid JSON object. Internal double-quotes escaped as \\\". "
        "No markdown fences. No comments. No trailing commas."
    )

    _PROJ_DIFFERENTIATORS = {
        "soc_auto":       ["SPL correlation searches","MITRE ATT&CK mapping","SOAR pipeline detail"],
        "vuln_scanner":   ["EPSS context","FIRST.org API mention","CVSS severity classification","remediation tracking evidence"],
        "phishing_osint": ["typosquatting detection detail","multi-API cross-referencing (VirusTotal, AbuseIPDB, URLScan.io)","WHOIS/DNS/SSL analysis detail"],
    }
    active_diffs = []
    for pk in set([p1_key, p2_key]):
        active_diffs.extend(_PROJ_DIFFERENTIATORS.get(pk, []))
    diff_instruction = (
        f"NEVER drop from project bullets: {', '.join(active_diffs)}.\n"
        "These differentiators are what make a fresher resume stand out — keep them even if longer.\n"
        "IMPORTANT: Only include technical details that belong to each specific project.\n"
        if active_diffs else ""
    )

    gap_ctx = gap.get("gap_instruction", "")

    user = f"""JOB:
  Title:   {job['job_title']}
  Company: {job['company']}
  Domain:  {job['domain']}
  Summary: {job['summary']}
  Skills:  {job['skills']}
{co_ctx}{kw_hint}
{lens_ctx}
	{gap_ctx}
	{research_ctx}
	ROLE-MARKET PROJECT SIGNALS:
	Project angles recruiters expect: {market_project_angles}
	Project keywords from Reddit/GitHub/web: {market_project_keywords}
	
	SINGLE-PAGE PREFERENCE: Keep bullets concise (prefer under 200 chars).
{diff_instruction}
	PROJECT METRIC BAN:
	- Do not write project metrics, impact claims, percentages, time savings, MTTR/MTTD, SLA deadlines, or reduced/improved/increased claims.
	- Project bullets should show what was built, configured, analyzed, documented, or mapped.
	PROJECT EVIDENCE RULE:
	- Generate project bullets from the verified project evidence below, not from generic role assumptions.
	- You may reframe wording for this JD, but do not add tools, integrations, APIs, certifications, platforms, or results missing from the evidence.
	- Keep project titles exactly as provided.
	PLACEHOLDER BAN:
	- Never write placeholder names such as P1_TECH, P2_TECH, P1 project, or P2 project in the final bullets.
	- Do not include raw SPL query syntax or MITRE technique IDs; write those concepts in plain English.

	PROJECT 1 VERIFIED EVIDENCE:
	Key: {p1_key}
	Title: {p1['title']}
	Allowed tools: {', '.join(p1_tools)}
	Tailoring guidance: {p1_guidance}
	{project_evidence.get(p1_key, '')[:2200]}

	PROJECT 2 VERIFIED EVIDENCE:
	Key: {p2_key}
	Title: {p2['title']}
	Allowed tools: {', '.join(p2_tools)}
	Tailoring guidance: {p2_guidance}
	{project_evidence.get(p2_key, '')[:2200]}
	
	Return JSON with EXACTLY 10 keys:
	{{
	  "P1_TITLE": "{p1['title']}",
	  "P1_TECH":  "{', '.join(p1_tools)}",
	  "P1_B1": "JD-tailored project bullet from P1 evidence only",
	  "P1_B2": "JD-tailored project bullet from P1 evidence only",
	  "P1_B3": "JD-tailored project bullet from P1 evidence only",
	  "P2_TITLE": "{p2['title']}",
	  "P2_TECH":  "{', '.join(p2_tools)}",
	  "P2_B1": "JD-tailored project bullet from P2 evidence only",
	  "P2_B2": "JD-tailored project bullet from P2 evidence only",
	  "P2_B3": "JD-tailored project bullet from P2 evidence only"
	}}
	Rules: 'and' not '&' | no project metrics or impact claims | escape internal quotes | each project uses ONLY its own technical details"""

    p1_seed_bullets = get_project_bullets(p1_key, job.get("domain", "General"))
    p2_seed_bullets = get_project_bullets(p2_key, job.get("domain", "General"))
    fallback_content = {
        "P1_TITLE": p1["title"], "P1_TECH": ", ".join(p1_tools),
        "P1_B1": p1_seed_bullets[0], "P1_B2": p1_seed_bullets[1], "P1_B3": p1_seed_bullets[2],
        "P2_TITLE": p2["title"], "P2_TECH": ", ".join(p2_tools),
        "P2_B1": p2_seed_bullets[0], "P2_B2": p2_seed_bullets[1], "P2_B3": p2_seed_bullets[2],
    }

    try:
        raw = _call_groq(system, user, GROQ_GEN_MODEL)
        raw = _repair_json(raw)
        try:
            content = json.loads(raw)
        except json.JSONDecodeError as exc:
            logger.warning("  JSON parse failed (%s) — repairing...", exc)
            fixed = re.sub(
                r'("(?:P[12]_(?:TITLE|TECH|B\d))":\s*)"(.*?)"(?=\s*[,}])',
                lambda m: m.group(1)+'"'+m.group(2).replace('"', '\\"')+'"',
                raw, flags=re.DOTALL
            )
            content = json.loads(fixed)
    except Exception as exc:
        logger.warning("  Project bullet generation failed — using deterministic fallback: %s", exc)
        content = dict(fallback_content)

    expected = ["P1_TITLE","P1_TECH","P1_B1","P1_B2","P1_B3","P2_TITLE","P2_TECH","P2_B1","P2_B2","P2_B3"]
    missing  = [k for k in expected if k not in content]
    if missing:
        logger.warning("  Missing keys from generation (%s) — filling from fallback", missing)
        for key, value in fallback_content.items():
            content.setdefault(key, value)

    # Skills: AI strategy first, deterministic profile as a safety fallback
    base_skills = (tailoring_strategy or {}).get("skills") or compute_skills(job["domain"])
    content.update(dynamic_skills_augment(base_skills, _keywords_with_market(jd_keywords, role_market)))

    # Synonym expansion (Feature 2)
    for k in ["P1_B1","P1_B2","P1_B3","P2_B1","P2_B2","P2_B3"]:
        if content.get(k):
            content[k] = apply_synonyms(content[k])

    # Project bullet sanitisation (strips unverifiable metrics + purpose clauses)
    content = sanitize_project_bullets(content)
    content = strip_cross_project_terms(content, p1_key, p2_key)

    # Inject pre-generated Amazon bullets and sanitize them
    content.update(amazon_bullets)
    content = sanitize_amazon_bullets(content, job, experience_research)

    return content


# ─────────────────────────────────────────────────────────────────────────────
# VALIDATION
# ─────────────────────────────────────────────────────────────────────────────
def _normalize_validation_output(data: dict) -> dict:
    return {
        "ats_score":        str(data.get("ats_score", "N/A")),
        "missing_keywords": str(data.get("missing_keywords", "")),
        "improvements":     str(data.get("improvements", "")),
        "github_insight":   str(data.get("github_insight", "")),
    }


def validate_resume(content: dict, job: dict, github_notes: str, mode: str) -> dict:
    EMPTY = {"ats_score":"skipped","missing_keywords":"","improvements":"","github_insight":""}
    if mode == "lenient":
        logger.info("  Validation: lenient — skipped")
        return EMPTY
    bullets = " | ".join(filter(None, [
        content.get("AMZ_B1",""), content.get("AMZ_B2",""),
        content.get("AMZ_B3",""), content.get("AMZ_B4",""),
        content.get("P1_B1",""),  content.get("P1_B2",""),
        content.get("P2_B1",""),  content.get("P2_B2",""),
    ]))
    if mode == "normal":
        prompt = (
            f"Job: {job.get('job_title','')} | JD keywords: {job.get('skills','')[:200]}\n"
            f"Bullets: {bullets[:500]}\nATS review for 0-2yr cybersecurity candidate.\n"
            "Return raw JSON: {\"ats_score\":<1-10>,\"missing_keywords\":\"<max 6>\"}"
        )
        try:
            raw  = _call_groq("Return only valid JSON, no markdown.", prompt, GROQ_VAL_MODEL, max_tokens=150)
            data = json.loads(_repair_json(raw))
            data = _normalize_validation_output(data)
            logger.info("  ATS=%s missing=%s", data.get("ats_score"), str(data.get("missing_keywords",""))[:50])
            return data
        except Exception as exc:
            logger.warning("  Validation failed: %s", exc)
            return EMPTY
    gh_sec = (f"\nSimilar GitHub projects:\n{github_notes[:500]}\n" if github_notes else "")
    prompt = (
        f"Job: {job.get('job_title','')} | Domain: {job.get('domain','')}\n"
        f"JD: {job.get('skills','')[:250]}\nBullets: {bullets[:600]}\n{gh_sec}"
        "Return raw JSON: {\"ats_score\":<1-10>,\"missing_keywords\":\"<max 8>\","
        "\"improvements\":\"<2 fixes>\",\"github_insight\":\"<1 thing>\"}"
    )
    try:
        raw  = _call_groq("Strict ATS reviewer. Return only valid JSON.", prompt, GROQ_VAL_MODEL, max_tokens=300)
        data = json.loads(_repair_json(raw))
        data = _normalize_validation_output(data)
        logger.info("  ATS=%s", data.get("ats_score"))
        return data
    except Exception as exc:
        logger.warning("  Validation failed: %s", exc)
        return EMPTY


# ─────────────────────────────────────────────────────────────────────────────
# DOCX FILL
# ─────────────────────────────────────────────────────────────────────────────
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _replace_in_para(para, placeholder: str, replacement: str) -> bool:
    all_t = para._p.findall(f".//{{{W_NS}}}t")
    for t in all_t:
        if t.text and placeholder in t.text:
            t.text = t.text.replace(placeholder, replacement)
            if t.text and (t.text[0] == " " or t.text[-1] == " "):
                t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            return True
    full = "".join(t.text or "" for t in all_t)
    if placeholder not in full:
        return False
    new_text = full.replace(placeholder, replacement)
    if all_t:
        all_t[0].text = new_text
        if new_text and (new_text[0] == " " or new_text[-1] == " "):
            all_t[0].set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        for t in all_t[1:]:
            t.text = ""
    return True


def fill_template(content: dict) -> bytes:
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError("resume_template.docx not found.")
    doc = Document(str(TEMPLATE_PATH))
    replacements = {f"[[{k}]]": v for k, v in content.items()}
    for para in doc.paragraphs:
        full = "".join(t.text or "" for t in para._p.findall(f".//{{{W_NS}}}t"))
        for ph, val in replacements.items():
            if ph in full:
                _replace_in_para(para, ph, val)
                full = full.replace(ph, val)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ─────────────────────────────────────────────────────────────────────────────
# PDF GENERATION
# ─────────────────────────────────────────────────────────────────────────────
def generate_pdf(docx_bytes: bytes) -> bytes:
    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = os.path.join(tmpdir, "resume.docx")
        with open(docx_path, "wb") as f:
            f.write(docx_bytes)
        result = subprocess.run(
            ["libreoffice","--headless","--convert-to","pdf:writer_pdf_Export","--outdir",tmpdir,docx_path],
            capture_output=True, text=True, timeout=60,
        )
        if result.returncode != 0:
            raise RuntimeError(f"LibreOffice: {result.stderr[:200]}")
        pdf_path = os.path.join(tmpdir, "resume.pdf")
        if not os.path.exists(pdf_path):
            raise FileNotFoundError("LibreOffice did not produce resume.pdf")
        with open(pdf_path, "rb") as f:
            return f.read()


# ─────────────────────────────────────────────────────────────────────────────
# SINGLE-PAGE ENFORCEMENT  (Feature H)
# ─────────────────────────────────────────────────────────────────────────────
def _count_pdf_pages(pdf_bytes: bytes) -> int:
    try:
        import pikepdf
        return len(pikepdf.open(io.BytesIO(pdf_bytes)).pages)
    except Exception as e:
        logger.warning("Page count failed: %s", e)
        return 999


def _score_bullet_relevancy(bullet_text: str, ranked_keywords: list) -> int:
    if not bullet_text or not ranked_keywords:
        return 0
    lower = bullet_text.lower()
    return sum(
        1 for kw in ranked_keywords[:10]
        if re.search(rf"(?<!\w){re.escape(kw)}(?!\w)", lower, re.IGNORECASE)
    )


def _shorten_bullet_deterministic(bullet_text: str, target_chars: int = 160) -> str:
    bullet = re.sub(r"\s+", " ", bullet_text or "").strip()
    if not bullet or len(bullet) <= target_chars:
        return bullet_text

    compact = re.sub(r"\s*\([^)]{12,90}\)", "", bullet)
    compact = _strip_purpose_clause(compact)
    if 45 <= len(compact) <= target_chars:
        return compact.rstrip(" ;,") + ("" if compact.rstrip().endswith((".", "!", "?")) else ".")

    candidates = []
    for sep in ("; ", ". ", ", and ", ", with ", ", for ", ", "):
        idx = compact.rfind(sep, 0, target_chars + 1)
        if idx >= 75:
            candidates.append(compact[:idx])
    if candidates:
        shortened = max(candidates, key=len)
    else:
        shortened = compact[:target_chars].rsplit(" ", 1)[0]
    shortened = shortened.rstrip(" ;,.-")
    if len(shortened) < 45:
        return bullet_text
    if shortened[-1] not in ".!?":
        shortened += "."
    logger.info("    Deterministic shorten %d->%d chars", len(bullet_text), len(shortened))
    return shortened


def _shorten_bullet_llm(bullet_text: str, target_chars: int = 160) -> str:
    if not bullet_text or len(bullet_text) <= target_chars:
        return bullet_text
    if not USE_LLM_SHORTENING:
        return _shorten_bullet_deterministic(bullet_text, target_chars)
    system = (
        f"You are a resume bullet editor. Shorten the bullet to under {target_chars} characters. "
        "PRESERVE: EPSS context, MITRE mapping, SOAR detail, FIRST.org mention. "
        "Use 'and' not '&'. Return ONLY the shortened bullet, no quotes, no explanation."
    )
    user = f"Shorten this resume bullet to ~{target_chars} chars:\n{bullet_text}"
    try:
        result = _call_groq(system, user, GROQ_GEN_MODEL, max_tokens=250)
        result = result.strip().strip('"')
        if len(result) > 20:
            logger.info("    Shortened %d→%d chars", len(bullet_text), len(result))
            return result
    except Exception as exc:
        logger.warning("    Bullet shortening failed: %s", exc)
    return _shorten_bullet_deterministic(bullet_text, target_chars)


def _trim_skills_line(skills_value: str, max_items: int = 4) -> str:
    if not skills_value:
        return skills_value
    items = [x.strip() for x in skills_value.split(",") if x.strip()]
    if len(items) <= max_items:
        return skills_value
    trimmed = ", ".join(items[:max_items])
    logger.info("    Skills trimmed: %d→%d items", len(items), max_items)
    return trimmed


def _reduce_paragraph_spacing(docx_bytes: bytes) -> bytes:
    doc = Document(io.BytesIO(docx_bytes))
    from docx.shared import Pt
    for para in doc.paragraphs:
        pf   = para.paragraph_format
        text = para.text.strip()
        if not text:
            pf.space_before = Pt(0)
            pf.space_after  = Pt(0)
            continue
        if para.style and para.style.name and "Heading" in para.style.name:
            pf.space_before = Pt(2)
            pf.space_after  = Pt(0)
        else:
            if pf.space_before is None or pf.space_before > Pt(2):
                pf.space_before = Pt(1)
            if pf.space_after is None or pf.space_after > Pt(2):
                pf.space_after = Pt(0)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


_SECTION_TITLES = {"education","work experience","projects","technical skills","certifications"}


def _is_section_header(para) -> bool:
    return para.text.strip().lower() in _SECTION_TITLES


def _is_skill_row(para) -> bool:
    text = para.text.strip()
    if not text or len(text) < 5:
        return False
    if ":" in text:
        label = text.split(":")[0].strip()
        if 3 <= len(label) <= 30:
            return True
    return False


def _expand_spacing_to_fill_page(docx_bytes: bytes, extra_pts: float) -> bytes:
    from docx.shared import Pt
    doc             = Document(io.BytesIO(docx_bytes))
    section_headers = []
    skill_rows      = []
    bullet_rows     = []
    separator_rows  = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            separator_rows.append(para)
        elif _is_section_header(para):
            section_headers.append(para)
        elif _is_skill_row(para):
            skill_rows.append(para)
        elif para.style and para.style.name == "List Paragraph":
            bullet_rows.append(para)
    def _get_pts(val):
        if val is None or val == 0:
            return 0.0
        return val / 12700.0
    n_h = max(len(section_headers), 1)
    n_s = max(len(skill_rows), 1)
    n_b = max(len(bullet_rows), 1)
    n_e = max(len(separator_rows), 1)
    for para in section_headers:
        pf = para.paragraph_format
        pf.space_before = Pt(_get_pts(pf.space_before) + extra_pts * 0.40 / n_h)
    for para in skill_rows:
        pf = para.paragraph_format
        share = extra_pts * 0.20 / n_s
        pf.space_before = Pt(_get_pts(pf.space_before) + share * 0.5)
        pf.space_after  = Pt(_get_pts(pf.space_after)  + share * 0.5)
    for para in bullet_rows:
        pf = para.paragraph_format
        pf.space_after = Pt(_get_pts(pf.space_after) + extra_pts * 0.20 / n_b)
    for para in separator_rows:
        pf = para.paragraph_format
        share = extra_pts * 0.20 / n_e
        pf.space_before = Pt(_get_pts(pf.space_before) + share * 0.5)
        pf.space_after  = Pt(_get_pts(pf.space_after)  + share * 0.5)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _fill_page(docx_bytes: bytes) -> tuple[bytes, bytes]:
    pdf_bytes = generate_pdf(docx_bytes)
    pages = _count_pdf_pages(pdf_bytes)
    if pages != 1:
        return docx_bytes, pdf_bytes
    trial_docx = _expand_spacing_to_fill_page(docx_bytes, 5.0)
    trial_pdf  = generate_pdf(trial_docx)
    if _count_pdf_pages(trial_pdf) > 1:
        logger.info("  Page fill: page already near-full, no expansion possible")
        return docx_bytes, pdf_bytes
    lo, hi = 0.0, 200.0
    best_docx, best_pdf = docx_bytes, pdf_bytes
    trial_docx = _expand_spacing_to_fill_page(docx_bytes, hi)
    trial_pdf  = generate_pdf(trial_docx)
    if _count_pdf_pages(trial_pdf) <= 1:
        logger.info("  Page fill: distributed 200.0pt (maximum), still fits")
        return trial_docx, trial_pdf
    logger.info("  Page fill: binary searching optimal spacing (0-200pt)...")
    for _ in range(10):
        mid = (lo + hi) / 2
        trial_docx = _expand_spacing_to_fill_page(docx_bytes, mid)
        trial_pdf  = generate_pdf(trial_docx)
        if _count_pdf_pages(trial_pdf) <= 1:
            lo, best_docx, best_pdf = mid, trial_docx, trial_pdf
        else:
            hi = mid
    logger.info("  Page fill complete: distributed %.1fpt of spacing", lo)
    return best_docx, best_pdf


def _generate_and_check(working: dict, reduce_spacing: bool = False,
                        fill_page: bool = False) -> tuple[bytes, bytes, int]:
    docx_bytes = fill_template(working)
    if reduce_spacing:
        docx_bytes = _reduce_paragraph_spacing(docx_bytes)
    pdf_bytes = generate_pdf(docx_bytes)
    pages     = _count_pdf_pages(pdf_bytes)
    if fill_page and pages == 1:
        docx_bytes, pdf_bytes = _fill_page(docx_bytes)
        pages = _count_pdf_pages(pdf_bytes)
    return docx_bytes, pdf_bytes, pages


def enforce_single_page(content: dict, job: dict,
                        jd_keywords: dict | None = None) -> tuple[bytes, bytes, str]:
    ranked   = (jd_keywords or {}).get("ranked", [])
    trim_log = []
    working  = dict(content)

    docx_bytes, pdf_bytes, pages = _generate_and_check(working)
    if pages <= 1:
        logger.info("  Single page OK — filling page")
        docx_bytes, pdf_bytes = _fill_page(docx_bytes)
        return docx_bytes, pdf_bytes, "page-filled"

    logger.info("  %d pages detected — enforcing single page", pages)

    # Tier 0: spacing reduction
    logger.info("  Tier 0: reducing paragraph spacing")
    docx_bytes, pdf_bytes, pages = _generate_and_check(working, reduce_spacing=True)
    trim_log.append("reduced-spacing")
    if pages <= 1:
        docx_bytes, pdf_bytes = _fill_page(docx_bytes)
        return docx_bytes, pdf_bytes, "; ".join(trim_log) + "; page-filled"

    # Tier 1: shorten long bullets (>200 chars)
    all_bullet_keys = ["AMZ_B1","AMZ_B2","AMZ_B3","AMZ_B4","P1_B1","P1_B2","P1_B3","P2_B1","P2_B2","P2_B3"]
    long_bullets    = sorted(
        [(k, len(working.get(k,""))) for k in all_bullet_keys if len(working.get(k,"")) > 200],
        key=lambda x: x[1], reverse=True,
    )
    if long_bullets:
        for key, length in long_bullets:
            working[key] = _shorten_bullet_llm(working[key], target_chars=150)
            trim_log.append(f"shortened {key} ({length}→{len(working[key])})")
        docx_bytes, pdf_bytes, pages = _generate_and_check(working, reduce_spacing=True)
        if pages <= 1:
            docx_bytes, pdf_bytes = _fill_page(docx_bytes)
            return docx_bytes, pdf_bytes, "; ".join(trim_log) + "; page-filled"

    # Tier 1.5: aggressive shortening (>150 chars)
    still_long = sorted(
        [(k, len(working.get(k,""))) for k in all_bullet_keys if len(working.get(k,"")) > 150],
        key=lambda x: x[1], reverse=True,
    )
    if still_long:
        for key, length in still_long:
            working[key] = _shorten_bullet_llm(working[key], target_chars=130)
            trim_log.append(f"aggressively shortened {key} ({length}→{len(working[key])})")
        docx_bytes, pdf_bytes, pages = _generate_and_check(working, reduce_spacing=True)
        if pages <= 1:
            docx_bytes, pdf_bytes = _fill_page(docx_bytes)
            return docx_bytes, pdf_bytes, "; ".join(trim_log) + "; page-filled"

    # Tier 2: remove least-relevant project bullet
    removable = [k for k in ["P1_B1","P1_B2","P1_B3","P2_B1","P2_B2","P2_B3"]
                 if working.get(k,"").strip() and working[k].strip() != " "]
    if removable:
        scored = sorted([(k, _score_bullet_relevancy(working[k], ranked)) for k in removable], key=lambda x: x[1])
        for key, score in scored:
            working[key] = " "
            trim_log.append(f"removed {key} (score={score})")
            docx_bytes, pdf_bytes, pages = _generate_and_check(working, reduce_spacing=True)
            if pages <= 1:
                docx_bytes, pdf_bytes = _fill_page(docx_bytes)
                return docx_bytes, pdf_bytes, "; ".join(trim_log) + "; page-filled"

    # Tier 3: trim skills
    for sk_key in ["SK_V5","SK_V4","SK_V3","SK_V2","SK_V1"]:
        original = working.get(sk_key, "")
        if original and len(original.split(",")) > 3:
            working[sk_key] = _trim_skills_line(original, max_items=3)
            trim_log.append(f"trimmed {sk_key}")
            docx_bytes, pdf_bytes, pages = _generate_and_check(working, reduce_spacing=True)
            if pages <= 1:
                docx_bytes, pdf_bytes = _fill_page(docx_bytes)
                return docx_bytes, pdf_bytes, "; ".join(trim_log) + "; page-filled"

    # Tier 4: shorten Amazon bullets (last resort)
    amz_bullets = sorted(
        [(k, len(working.get(k,""))) for k in AMAZON_KEYS if working.get(k,"").strip() and working[k].strip() != " "],
        key=lambda x: x[1], reverse=True,
    )
    for key, length in amz_bullets:
        if length > 80:
            working[key] = _shorten_bullet_llm(working[key], target_chars=100)
            trim_log.append(f"shortened {key} ({length}→{len(working[key])})")
            docx_bytes, pdf_bytes, pages = _generate_and_check(working, reduce_spacing=True)
            if pages <= 1:
                docx_bytes, pdf_bytes = _fill_page(docx_bytes)
                return docx_bytes, pdf_bytes, "; ".join(trim_log) + "; page-filled"

    logger.warning("  All tiers exhausted — could not achieve single page")
    docx_bytes, pdf_bytes, _ = _generate_and_check(working, reduce_spacing=True)
    return docx_bytes, pdf_bytes, "; ".join(trim_log) + "; overflow-unresolved"


# ─────────────────────────────────────────────────────────────────────────────
# GITHUB STORAGE + URL SHORTENING
# ─────────────────────────────────────────────────────────────────────────────
def _safe(s: str, n: int = 35) -> str:
    return re.sub(r"[^A-Za-z0-9_-]", "_", s)[:n]


def _github_commit(filename: str, file_bytes: bytes, message: str) -> str:
    path    = f"{RESUMES_FOLDER}/{filename}"
    api_url = f"https://api.github.com/repos/{GITHUB_REPOSITORY}/contents/{path}"
    headers = {
        "Authorization": f"Bearer {GITHUB_TOKEN}",
        "Accept": "application/vnd.github+json",
        "X-GitHub-Api-Version": "2022-11-28",
    }
    sha      = None
    existing = requests.get(api_url, headers=headers, timeout=10)
    if existing.status_code == 200:
        sha = existing.json().get("sha")
    payload = {"message": message, "content": base64.b64encode(file_bytes).decode(), "branch": GITHUB_BRANCH}
    if sha:
        payload["sha"] = sha
    resp = requests.put(api_url, json=payload, headers=headers, timeout=30)
    resp.raise_for_status()
    return f"https://raw.githubusercontent.com/{GITHUB_REPOSITORY}/{GITHUB_BRANCH}/{path}"


def upload_to_github(docx_bytes: bytes, pdf_bytes: bytes, job: dict) -> tuple[str, str]:
    base = f"Resume_{_safe(job['job_title'])}_{_safe(job['company'])}"
    msg  = f"Resume: {job['job_title']} @ {job['company']}"
    return (
        _github_commit(f"{base}.docx", docx_bytes, msg),
        _github_commit(f"{base}.pdf",  pdf_bytes,  msg),
    )


def shorten_url(long_url: str) -> str:
    try:
        resp = requests.get(
            f"https://tinyurl.com/api-create.php?url={requests.utils.quote(long_url)}",
            timeout=8,
        )
        if resp.status_code == 200 and resp.text.startswith("https://tinyurl.com"):
            return resp.text.strip()
    except Exception:
        pass
    return long_url


# ─────────────────────────────────────────────────────────────────────────────
# SHEETS HELPERS
# ─────────────────────────────────────────────────────────────────────────────
def _get_creds() -> Credentials:
    j = os.environ.get("GOOGLE_CREDS_JSON", "")
    if not j:
        raise EnvironmentError("GOOGLE_CREDS_JSON not set.")
    return Credentials.from_service_account_info(json.loads(j), scopes=SCOPES)


def ensure_column(ws, name: str) -> int:
    headers = ws.row_values(1)
    if name not in headers:
        idx = len(headers) + 1
        ws.update_cell(1, idx, name)
        headers.append(name)
        logger.info("Added column '%s' at %d.", name, idx)
        return idx
    return headers.index(name) + 1


def get_pending_jobs(ws, doc_col: int) -> list[dict]:
    rows = ws.get_all_values()
    if len(rows) < 2:
        return []
    headers = rows[0]
    col = {h: i for i, h in enumerate(headers)}
    def _get(row, key):
        i = col.get(key)
        return row[i].strip() if i is not None and i < len(row) else ""
    pending = []
    for row_num, row in enumerate(rows[1:], start=2):
        if _get(row, "status").lower() == "new" and not (row[doc_col-1].strip() if doc_col-1 < len(row) else ""):
            pending.append({
                "row_num":   row_num,
                "job_title": _get(row, "job_title") or "Cybersecurity Role",
                "company":   _get(row, "company")   or "Unknown",
                "domain":    _get(row, "domain")     or "General",
                "summary":   _get(row, "summary"),
                "skills":    _get(row, "skills_required"),
            })
    return pending


# ─────────────────────────────────────────────────────────────────────────────
# MAIN
# [CHANGE H] Three-phase structure:
#   Phase 1 — extract_keywords() for all jobs (fast, sklearn or Groq)
#   Phase 2 — collect_experience_research() in parallel via ThreadPoolExecutor
#   Phase 3 — sequential generate + validate + upload per job
# ─────────────────────────────────────────────────────────────────────────────
def main():
    global _groq_call_count, _groq_consecutive_429s, _groq_cooldown_until, _groq_last_call_ts
    _groq_call_count = 0   # reset counter at start of run
    _groq_consecutive_429s = 0
    _groq_cooldown_until = 0.0
    _groq_last_call_ts = 0.0

    logger.info("=" * 60)
    logger.info("Resume Tailor — Research Framework Edition (validation=%s)", VALIDATION_MODE)
    logger.info(
        "Groq controls: model=%s retries=%d min_interval=%.1fs market_llm=%s recruiter=%s llm_shortening=%s max_jobs=%d",
        GROQ_GEN_MODEL, GROQ_MAX_RETRIES, GROQ_MIN_INTERVAL_SECONDS,
        ROLE_MARKET_LLM_SUMMARY, RECRUITER_SIMULATION, USE_LLM_SHORTENING, MAX_JOBS_PER_RUN,
    )
    logger.info("=" * 60)

    for name, val in [("GROQ_API_KEY",GROQ_API_KEY),("GITHUB_TOKEN",GITHUB_TOKEN),("GITHUB_REPOSITORY",GITHUB_REPOSITORY)]:
        if not val:
            logger.error("%s not set.", name)
            sys.exit(1)
    if not TEMPLATE_PATH.exists():
        logger.error("resume_template.docx not found.")
        sys.exit(1)

    creds = _get_creds()
    gc    = gspread.authorize(creds)
    ws    = gc.open(SHEET_NAME).sheet1
    logger.info("Connected to Sheets.")

    doc_col      = ensure_column(ws, "resume_doc_link")
    pdf_col      = ensure_column(ws, "resume_pdf_link")
    val_col      = ensure_column(ws, "validation_notes")
    cov_col      = ensure_column(ws, "keyword_coverage")
    den_col      = ensure_column(ws, "keyword_density")
    sk_col       = ensure_column(ws, "total_skills_count")
    cred_col     = ensure_column(ws, "credibility")
    stuff_col    = ensure_column(ws, "stuffing_suspicion")
    hire_col     = ensure_column(ws, "hireability")
    research_col = ensure_column(ws, "research_sources")

    pending = get_pending_jobs(ws, doc_col)
    if not pending:
        logger.info("No New jobs with empty resume_doc_link.")
        sys.exit(0)

    logger.info("Found %d pending. Processing up to %d.", len(pending), MAX_JOBS_PER_RUN)
    pending = pending[:MAX_JOBS_PER_RUN]

    # ── Phase 1: Keyword extraction for all jobs (fast — sklearn or single Groq call each)
    logger.info("Phase 1: Extracting keywords for %d jobs...", len(pending))
    all_keywords: dict[int, dict] = {}
    for job in pending:
        # [CHANGE A] Weighted JD: title 3x, skills 2x, summary 1x
        weighted_jd = (
            f"{job.get('job_title', '')} " * 3 +
            f"{job.get('skills', '')} " * 2 +
            job.get('summary', '')
        )
        all_keywords[job['row_num']] = extract_keywords(weighted_jd)
        logger.info(
            "  [%s] top keywords: %s",
            job['job_title'][:30],
            all_keywords[job['row_num']].get('ranked', [])[:3],
        )

    # ── Phase 2: Parallel research (HTTP-bound — safe to parallelise)
    logger.info("Phase 2: Collecting experience + role-market research in parallel (max_workers=3)...")
    all_research: dict[int, dict] = {}
    all_market_intel: dict[int, dict] = {}
    with ThreadPoolExecutor(max_workers=3) as executor:
        research_futures = {
            executor.submit(
                collect_experience_research,
                job,
                all_keywords[job['row_num']],
            ): job
            for job in pending
        }
        for future in as_completed(research_futures):
            job = research_futures[future]
            try:
                all_research[job['row_num']] = future.result()
                logger.info(
                    "  Research done: %s @ %s",
                    job['job_title'][:25], job['company'][:20],
                )
            except Exception as exc:
                logger.warning("  Research failed for %s: %s", job['job_title'][:25], exc)
                all_research[job['row_num']] = {"summary": "", "source_note": f"research-failed: {exc}"}
    with ThreadPoolExecutor(max_workers=3) as executor:
        market_futures = {
            executor.submit(
                collect_role_market_intel,
                job,
                all_keywords[job['row_num']],
            ): job
            for job in pending
        }
        for future in as_completed(market_futures):
            job = market_futures[future]
            try:
                all_market_intel[job['row_num']] = future.result()
                logger.info(
                    "  Market intel done: %s @ %s",
                    job['job_title'][:25], job['company'][:20],
                )
            except Exception as exc:
                logger.warning("  Market intel failed for %s: %s", job['job_title'][:25], exc)
                all_market_intel[job['row_num']] = {"summary": "", "source_note": f"market-failed: {exc}"}

    # ── Phase 3: Sequential generation + validation + upload
    logger.info("Phase 3: Generating and uploading resumes...")
    success = 0
    for i, job in enumerate(pending, 1):
        logger.info("-" * 50)
        logger.info("[%d/%d] %s @ %s  (domain: %s)", i, len(pending),
                    job["job_title"], job["company"], job["domain"])
        try:
            jd_text            = f"{job['skills']} {job['summary']} {job['job_title']}"
            jd_keywords        = all_keywords[job['row_num']]
            experience_research = all_research.get(job['row_num'], {})
            role_market_intel   = all_market_intel.get(job['row_num'], {})

            tailoring_strategy = build_tailoring_strategy(
                job, jd_keywords, experience_research, role_market_intel
            )
            p1_key, p2_key = tailoring_strategy.get("projects", ["soc_auto","vuln_scanner"])[:2]
            project_evidence = {
                p1_key: get_project_evidence(p1_key, job.get("domain", "General")),
                p2_key: get_project_evidence(p2_key, job.get("domain", "General")),
            }
            p1_tools       = select_tools(p1_key, jd_text)
            p2_tools       = select_tools(p2_key, jd_text)
            logger.info(
                "  Projects: %s + %s | P1 tools: %s | strategy=%s",
                p1_key, p2_key, p1_tools[:3], tailoring_strategy.get("source", "unknown"),
            )

            # GitHub research (strict mode only)
            github_notes = ""
            if VALIDATION_MODE == "strict":
                github_notes = research_github_projects(job["domain"], job["job_title"])

            # Company intel
            intel       = get_company_intel(job["company"])
            scraped_ctx = "" if intel else scrape_company(job["company"])

            # Generate content
            logger.info("  Generating content...")
            content = generate_content(
                job, p1_key, p2_key, intel, scraped_ctx,
                p1_tools, p2_tools, jd_keywords, experience_research,
                tailoring_strategy, project_evidence,
            )

            # Track keyword usage (Feature 3)
            track_keyword_usage(content, jd_keywords.get("ranked", []))

            # Validate
            if VALIDATION_MODE != "lenient":
                time.sleep(3)
            val_result = validate_resume(content, job, github_notes, VALIDATION_MODE)
            ats_score  = val_result.get("ats_score", "N/A")
            val_note   = (
                f"[{VALIDATION_MODE.upper()}] ATS:{ats_score}"
                + (f" | Missing:{val_result.get('missing_keywords','')}" if val_result.get("missing_keywords") else "")
                + (f" | Fix:{val_result.get('improvements','')}"         if val_result.get("improvements")     else "")
                + (f" | GitHub:{val_result.get('github_insight','')}"    if val_result.get("github_insight")   else "")
            )
            logger.info("  %s", val_note)

            # Metrics (Feature 5)
            metrics = compute_metrics(content, jd_keywords, ats_score)

            # Recruiter simulation (Feature 6)
            if RECRUITER_SIMULATION and VALIDATION_MODE != "lenient":
                time.sleep(2)
                rec_sim = recruiter_simulate(content, job)
            else:
                rec_sim = {"credibility":"skipped","stuffing_suspicion":"skipped","hireability":"skipped"}

            # Single-page enforcement + PDF generation (Feature H)
            logger.info("  Generating DOCX+PDF (single-page enforcement)...")
            docx_bytes, pdf_bytes, trim_log = enforce_single_page(content, job, jd_keywords)
            if trim_log and trim_log not in ("certs-p2-ok", ""):
                val_note += f" | Trimmed:{trim_log}"
            logger.info("  DOCX: %d bytes  PDF: %d bytes", len(docx_bytes), len(pdf_bytes))

            # Upload + shorten URLs
            doc_raw, pdf_raw = upload_to_github(docx_bytes, pdf_bytes, job)
            doc_url = shorten_url(doc_raw)
            pdf_url = shorten_url(pdf_raw)
            logger.info("  Doc: %s", doc_url)
            logger.info("  PDF: %s", pdf_url)

            # Write all columns to sheet
            ws.update_cell(job["row_num"], doc_col,      doc_url)
            ws.update_cell(job["row_num"], pdf_col,      pdf_url)
            ws.update_cell(job["row_num"], val_col,      val_note)
            ws.update_cell(job["row_num"], cov_col,      metrics["keyword_coverage"])
            ws.update_cell(job["row_num"], den_col,      metrics["keyword_density"])
            ws.update_cell(job["row_num"], sk_col,       metrics["total_skills_count"])
            ws.update_cell(job["row_num"], cred_col,     str(rec_sim.get("credibility", "")))
            ws.update_cell(job["row_num"], stuff_col,    str(rec_sim.get("stuffing_suspicion", "")))
            ws.update_cell(job["row_num"], hire_col,     str(rec_sim.get("hireability", "")))
            research_note = (
                f"experience: {experience_research.get('source_note', '')}; "
                f"market: {role_market_intel.get('source_note', '')}"
            )
            ws.update_cell(job["row_num"], research_col, research_note)
            logger.info("  ✓ Sheet updated.")

            success += 1
            time.sleep(4)

        except Exception as exc:
            logger.error("  ✗ Failed: %s", exc)
            continue

    # [CHANGE I] Log Groq call budget at end of run
    logger.info("=" * 60)
    logger.info("Done: %d/%d succeeded. | Groq API calls this run: %d", success, len(pending), _groq_call_count)
    logger.info("=" * 60)


if __name__ == "__main__":
    main()
